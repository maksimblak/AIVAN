# docx_export.py
from __future__ import annotations

import re
import tempfile
import uuid
from datetime import datetime
from html import unescape
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Mapping, Sequence

GARANT_WEB = "https://d.garant.ru"


# ---------- plumbing ----------

def _require_docx():
    """Lazy import python-docx with a clear error if missing."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, Cm  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx не установлен, DOCX-экспорт недоступен") from exc
    return Document, Pt, Cm


def _html_to_plain(text: str) -> str:
    """Very simple HTML→plain fallback (almost unused, kept for future)."""
    if not text:
        return ""
    clean = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    clean = re.sub(r"<[^>]+>", "", clean)
    return unescape(clean).strip()


def _temp_path(stem: str, suffix: str = ".docx") -> Path:
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", stem).strip("_") or "export"
    return Path(tempfile.gettempdir()) / f"{safe}_{uuid.uuid4().hex}{suffix}"


# ---------- styles & layout ----------

def _setup_page_and_styles(doc) -> None:
    """Configure A4 page, margins, and base styles."""
    from docx.shared import Pt, Cm  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    # A4 + margins
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)

    # Normal: TNR 12pt, 1.15 line spacing, 0pt before / 6pt after, justify
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        pf = normal.paragraph_format
        pf.line_spacing = 1.15
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except Exception:
        pass

    # Headings: bold, TNR, keep-with-next
    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)):
        try:
            st = doc.styles[name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(size)
            st.font.bold = True
            pf = st.paragraph_format
            pf.line_spacing = 1.15
            pf.space_before = Pt(6)
            pf.space_after = Pt(4)
            pf.keep_with_next = True
        except Exception:
            continue


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert a clickable hyperlink run into a paragraph."""
    if not url:
        paragraph.add_run(text)
        return
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception:  # pragma: no cover
        paragraph.add_run(f"{text} ({url})")
        return

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    # color blue
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000EE")
    r_pr.append(color)

    run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text or url
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _remove_paragraph(p) -> None:
    p._element.getparent().remove(p._element)


def _is_list_paragraph(p) -> bool:
    try:
        name = p.style.name if p.style else ""
        return "List" in name or "Список" in name
    except Exception:
        return False


def _tidy_document(doc) -> None:
    """Normalize spacing/indentation and remove stray tiny lines."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Pt, Cm  # type: ignore

    for p in list(doc.paragraphs):
        text = (p.text or "").strip().replace("\xa0", " ")

        # Drop useless standalone bullets/dashes or lone small numbers
        if text in {"—", "-", "•", "*"} or (text.isdigit() and len(text) <= 3):
            _remove_paragraph(p)
            continue

        # Headings: keep with next
        if p.style and str(p.style.name).startswith("Heading"):
            p.paragraph_format.keep_with_next = True
            continue

        # Normal paragraphs: justify, 1.15, 6pt after, first-line indent unless list
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        if pf.line_spacing is None:
            pf.line_spacing = 1.15
        if pf.space_after is None:
            pf.space_after = Pt(6)
        if not _is_list_paragraph(p):
            left_ind = getattr(pf, "left_indent", None)
            if not left_ind or getattr(left_ind, "pt", 0) == 0:
                if text:
                    pf.first_line_indent = Cm(1.25)


# ---------- numbering for case headings only ----------

def _autonumber_case_headings(doc) -> None:
    """
    Add simple 1., 2., 3. numbering to paragraphs with style 'Heading 3'
    (used for case titles only). Skip if the heading already starts with a number.
    """
    num = 0
    num_re = re.compile(r"^\s*\d+([.)]|\.\d+)\s|^\s*\d+\)\s")

    for p in doc.paragraphs:
        if not p.style or p.style.name != "Heading 3":
            continue
        text = (p.text or "").strip()
        if not text or num_re.match(text):
            continue
        num += 1
        p.text = f"{num}. {text}"


# ---------- simple HTML → docx renderer ----------

class _HTML2Docx(HTMLParser):
    def __init__(self, doc):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.paragraph = None
        self.bold = self.italic = self.underline = self.code = self.pre = False
        self.list_stack: list[str] = []
        self.href: str | None = None

    def _ensure_paragraph(self, style: str | None = None):
        if self.paragraph is None:
            self.paragraph = self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()
        return self.paragraph

    def _clean_data(self, data: str) -> str:
        data = data.replace("\xa0", " ").replace("\u00AD", "")
        if not self.pre:
            data = re.sub(r"\s+", " ", data)
        return data

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag == "p":
            self.paragraph = self.doc.add_paragraph()
        elif tag in {"h1", "h2", "h3"}:
            style = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3"}[tag]
            self.paragraph = self.doc.add_paragraph(style=style)
        elif tag == "div":
            return
        elif tag == "br":
            from docx.enum.text import WD_BREAK  # type: ignore
            self._ensure_paragraph().add_run().add_break(WD_BREAK.LINE)
        elif tag in {"b", "strong"}:
            self.bold = True
        elif tag in {"i", "em"}:
            self.italic = True
        elif tag == " u":
            self.underline = True
        elif tag == "code":
            self.code = True
        elif tag == "pre":
            self.pre = True
            self.paragraph = self.doc.add_paragraph()
        elif tag == "blockquote":
            self.paragraph = self.doc.add_paragraph()
            try:
                pf = self.paragraph.paragraph_format
                from docx.shared import Cm  # type: ignore
                pf.left_indent = Cm(1)
                pf.right_indent = Cm(1)
            except Exception:
                pass
        elif tag == "a":
            self.href = None
            for k, v in attrs:
                if k.lower() == "href":
                    self.href = make_url(v)
                    break
        elif tag == "ul":
            self.list_stack.append("ul")
        elif tag == "ol":
            self.list_stack.append("ol")
        elif tag == "li":
            style = "List Bullet" if (self.list_stack and self.list_stack[-1] == "ul") else "List Number"
            self.paragraph = self.doc.add_paragraph(style=style)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in {"p", "li", "pre", "h1", "h2", "h3", "blockquote"}:
            self.paragraph = None
        elif tag in {"ul", "ol"} and self.list_stack and self.list_stack[-1] == tag:
            self.list_stack.pop()
        elif tag in {"b", "strong"}:
            self.bold = False
        elif tag in {"i", "em"}:
            self.italic = False
        elif tag == "u":
            self.underline = False
        elif tag == "code":
            self.code = False
        elif tag == "pre":
            self.pre = False
        elif tag == "a":
            self.href = None

    def handle_data(self, data):
        if not data:
            return
        data = self._clean_data(data)
        if not data:
            return

        # swallow lone bullets / tiny numbers when not in <pre>
        if not self.pre:
            s = data.strip()
            if s in {"—", "-", "•", "*"} or (s.isdigit() and len(s) <= 3):
                return

        p = self._ensure_paragraph()
        if self.href:
            _add_hyperlink(p, data.strip(), self.href)
            return

        run = p.add_run(data)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline
        if self.code or self.pre:
            try:
                run.font.name = "Courier New"
            except Exception:
                pass


def _render_html(doc, html: str) -> None:
    if not html or not html.strip():
        return
    parser = _HTML2Docx(doc)
    parser.feed(html)
    parser.close()


# ---------- URL helpers ----------

_TOPIC_RE = re.compile(r"(?:/|#/)document/(\d+)", re.IGNORECASE)

def _topic_from_url(url: str) -> int | None:
    if not url:
        return None
    _TR = _TOPIC_RE.search(url)
    return int(_TR.group(1)) if _TR else None


def make_url(topic_or_href: Any) -> str:
    """
    Normalize a Garant topic id or href to a public URL like:
    https://d.garant.ru/document/{id}/
    """
    val = topic_or_href
    if isinstance(val, int):
        return f"{GARANT_WEB}/document/{val}/"

    s = str(val or "").strip()
    if not s:
        return f"{GARANT_WEB}/"

    # direct /document/123456 or #/document/123456
    m = _TOPIC_RE.search(s)
    if m:
        return f"{GARANT_WEB}/document/{m.group(1)}/"

    # any 5+ digit number fallback
    m = re.search(r"\d{5,}", s)
    if m:
        return f"{GARANT_WEB}/document/{m.group(0)}/"

    return f"{GARANT_WEB}/"


# ---------- main builder ----------

def build_pricing_docx(  # kept signature as in your code (typo fixed internally)
    *,
    summary_html: str,
    fragments: Sequence[Any],
    structured: Mapping[str, Any] | None = None,
    full_texts: Mapping[int, str] | None = None,
    file_stub: str | None = None,
) -> Path:
    """Create a nicely formatted DOCX with summary + cases + (optional) full texts."""
    Document, Pt, Cm = _require_docx()
    full_texts = dict(full_texts or {})

    doc = Document()
    _setup_page_and_styles(doc)

    # Title + timestamp
    doc.add_paragraph("Подборка судебной практики", style="Heading 1")
    doc.add_paragraph(datetime.now().strftime("%d.%m.%Y %H:%M"))

    # Summary
    doc.add_paragraph()  # spacer
    doc.add_paragraph("Краткий вывод", style="Heading 2")

    summary = None
    if structured and isinstance(structured.get("summary"), str):
        summary = structured["summary"]

    if isinstance(summary, str) and summary.strip():
        _render_html(doc, summary)
    elif summary_html and summary_html.strip():
        _render_html(doc, summary_html)
    else:
        p = doc.add_paragraph()
        p.add_run("нет данных").italic = True

    # Collect/normalize cases
    cases: list[Mapping[str, Any]] = []
    if structured and isinstance(structured.get("cases"), list):
        cases = [c for c in structured["cases"] if isinstance(c, Mapping)]
    elif fragments:
        for frag in fragments:
            match = getattr(frag, "match", None)
            meta = getattr(match, "metadata", {}) if match else {}
            if not isinstance(meta, Mapping):
                continue
            cases.append(
                {
                    "title": meta.get("title") or meta.get("name") or getattr(frag, "header", ""),
                    "case_number": meta.get("case_number") or meta.get("case"),
                    "url": meta.get("url") or meta.get("link"),
                    "facts": meta.get("summary") or getattr(frag, "excerpt", ""),
                    "holding": meta.get("decision_summary") or "",
                    "norms": meta.get("norms_summary") or "",
                    "topic": meta.get("topic"),
                    "entry": meta.get("entry"),
                    "fulltext_html": meta.get("fulltext_html") or "",
                }
            )

    # Cases section
    doc.add_paragraph()
    doc.add_paragraph("Подборка дел", style="Heading 2")

    if not cases:
        doc.add_paragraph().add_run("Нет данных по судебной практике").italic = True
    else:
        for idx, case in enumerate(cases, start=1):
            title = (str(case.get("title") or "")).strip()
            case_number = (str(case.get("case_number") or "")).strip()

            url = make_url(case.get("link") or case.get("url") or case.get("topic"))
            # store back normalized url if dict-like
            if isinstance(case, dict):
                case["url"] = url
                case["link"] = url

            facts = (str(case.get("facts") or "")).strip()
            holding = (str(case.get("holding") or "")).strip()

            norms_raw = case.get("norms")
            if isinstance(norms_raw, (list, tuple)):
                norms_text = ", ".join(str(it).strip() for it in norms_raw if str(it).strip())
            else:
                norms_text = (str(norms_raw or "")).strip()

            topic = case.get("topic")
            try:
                if topic is None:
                    t = _topic_from_url(url)
                    topic = int(t) if t is not None else None
                else:
                    topic = int(topic)
            except Exception:
                topic = None

            # --- case title as Heading 3 (numbered later) ---
            doc.add_paragraph(title or f"Дело {idx}", style="Heading 3")

            # case number
            if case_number:
                p = doc.add_paragraph()
                p.add_run("Номер дела: ").bold = True
                p.add_run(case_number)

            # link
            if url:
                p = doc.add_paragraph()
                p.add_run("Ссылка: ").bold = True
                _add_hyperlink(p, url, url)

            # facts
            if facts:
                p = doc.add_paragraph()
                p.add_run("Выдержка: ").bold = True
                _render_html(doc, facts)

            # holding
            if holding:
                p = doc.add_paragraph()
                p.add_run("Вывод суда: ").bold = True
                _render_html(doc, holding)

            # norms
            if norms_text:
                p = doc.add_paragraph()
                p.add_run("Нормы: ").bold = True
                for line in [s.strip() for s in norms_text.splitlines() if s.strip()]:
                    doc.add_paragraph(line, style="List Bullet")

            # full text
            body = (str(case.get("fulltext_html") or "")).strip()
            if not body and isinstance(topic, int):
                body = full_texts.get(topic) or ""
            if body:
                p = doc.add_paragraph()
                p.add_run("Полный текст:").bold = True
                _render_html(doc, body)
            else:
                p = doc.add_paragraph()
                p.add_run("Полный текст недоступен через API, используйте ссылку выше.").italic = True

            # spacer between cases
            doc.add_paragraph()

    # finalize: number only case headings, tidy, save
    _autonumber_case_headings(doc)
    _tidy_document(doc)

    out = _temp_path(file_stub or "practice_fulltext")
    doc.save(str(out))
    return out
