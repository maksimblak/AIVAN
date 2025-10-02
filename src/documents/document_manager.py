"""
Менеджер документов - центральный класс для управления всеми модулями работы с документами
"""

from __future__ import annotations
import asyncio
import importlib.util
import json
import logging
from html import escape as html_escape
from datetime import datetime
from pathlib import Path
from typing import Any, List

from .anonymizer import DocumentAnonymizer
from .base import DocumentResult, DocumentStorage, ProcessingError
from .storage_backends import ArtifactUploader
from .document_chat import DocumentChat
from .ocr_converter import OCRConverter
from .risk_analyzer import RiskAnalyzer
from .summarizer import DocumentSummarizer
from .translator import DocumentTranslator

# Импорт утилит для безопасной HTML сборки
from src.core.safe_telegram import format_safe_html, split_html_for_telegram
from src.bot.ui_components import sanitize_telegram_html
from src.core.settings import AppSettings

logger = logging.getLogger(__name__)


class DocumentManager:
    """Центральный менеджер для работы с документами"""
    RISK_LEVEL_LABELS = {
        "low": "низкий",
        "medium": "средний",
        "high": "высокий",
        "critical": "критический",
    }


    def __init__(
        self,
        openai_service=None,
        storage_path: str = "data/documents",
        *,
        storage_quota_mb: int | None = None,
        storage_cleanup_hours: int | None = None,
        storage_cleanup_interval: float | None = None,
        artifact_uploader: ArtifactUploader | None = None,
        settings: AppSettings | None = None,
    ):
        if settings is None:
            from src.core.app_context import get_settings  # avoid circular import

            settings = get_settings()

        self.settings = settings
        self.openai_service = openai_service

        quota_mb = storage_quota_mb if storage_quota_mb is not None else settings.document_storage_quota_mb
        cleanup_hours_value = (
            storage_cleanup_hours if storage_cleanup_hours is not None else settings.document_cleanup_hours
        )
        cleanup_interval_value = (
            storage_cleanup_interval
            if storage_cleanup_interval is not None
            else settings.document_cleanup_interval_seconds
        )

        uploader = artifact_uploader or self._build_artifact_uploader(settings)

        self.storage = DocumentStorage(
            storage_path,
            max_user_quota_mb=quota_mb,
            cleanup_max_age_hours=cleanup_hours_value if cleanup_hours_value is not None else 24,
            cleanup_interval_seconds=cleanup_interval_value if cleanup_interval_value is not None else 3600.0,
            artifact_uploader=uploader,
        )

        # Инициализация модулей
        self.summarizer = DocumentSummarizer(openai_service, settings=settings)
        self.risk_analyzer = RiskAnalyzer(openai_service)
        self.document_chat = DocumentChat(openai_service, settings=settings)
        self.anonymizer = DocumentAnonymizer(settings=settings)
        self.translator = DocumentTranslator(openai_service, settings=settings)
        self.ocr_converter = OCRConverter(settings=settings)
    def _build_artifact_uploader(self, settings: AppSettings) -> ArtifactUploader | None:
        bucket = settings.documents_s3_bucket
        if not bucket:
            return None
        try:
            return S3ArtifactUploader(
                bucket=bucket,
                prefix=settings.documents_s3_prefix,
                region_name=settings.documents_s3_region,
                endpoint_url=settings.documents_s3_endpoint,
                public_base_url=settings.documents_s3_public_url,
                acl=settings.documents_s3_acl or None,
            )
        except RuntimeError as exc:
            logger.warning("S3 artifact uploader disabled: %s", exc)
            return None



        self._dependencies: dict[str, bool] = {
            'docx': self._module_available('docx'),
            'reportlab': self._module_available('reportlab.pdfgen'),
        }

        self.PROCESSOR_PARAM_WHITELIST: dict[str, set[str]] = {
            "summarize": {"detail_level", "language", "output_formats"},
            "analyze_risks": {"custom_criteria"},
            "chat": set(),
            "anonymize": {"anonymization_mode", "exclude_types", "custom_patterns"},
            "translate": {"source_lang", "target_lang", "output_formats"},
            "ocr": {"output_format"},
        }

    @staticmethod
    def _module_available(module: str) -> bool:
        return importlib.util.find_spec(module) is not None

    def _dependency_available(self, key: str) -> bool:
        return self._dependencies.get(key, True)

    def _dependency_notice(self, *, dependency: str, feature: str, format_name: str) -> dict[str, str]:
        message = f"{feature} requires {dependency} to be installed."
        logger.warning(message)
        return {"format": format_name, "error": message}

    def _localize_risk_level(self, level: str | None, *, uppercase: bool = False, capitalize: bool = False) -> str:
        if level is None:
            label = "неизвестный"
        else:
            key = str(level).lower()
            label = self.RISK_LEVEL_LABELS.get(key, str(level))
        if uppercase:
            return label.upper()
        if capitalize:
            return label.capitalize()
        return label



    async def process_document(
        self,
        user_id: int,
        file_content: bytes,
        original_name: str,
        mime_type: str,
        operation: str,
        **kwargs: Any,
    ) -> DocumentResult:
        """Обработать документ указанным способом."""

        try:
            document_info = await self.storage.save_document(
                user_id, file_content, original_name, mime_type
            )

            processor = self._get_processor(operation)
            if not processor:
                raise ProcessingError(f"Неизвестная операция: {operation}", "UNKNOWN_OPERATION")

            allowed_params = self.PROCESSOR_PARAM_WHITELIST.get(operation, set())
            safe_kwargs: dict[str, Any] = {}
            for key, value in kwargs.items():
                if key not in allowed_params or value is None:
                    continue
                if key == "exclude_types" and isinstance(value, list):
                    normalized = sorted({str(item).lower() for item in value if item})
                    safe_kwargs[key] = normalized
                elif key == "custom_patterns" and isinstance(value, list):
                    sanitized = []
                    for item in value:
                        if isinstance(item, str):
                            pattern = item.strip()
                            if pattern:
                                sanitized.append({"pattern": pattern})
                        elif isinstance(item, dict) and item.get("pattern"):
                            pattern = str(item["pattern"]).strip()
                            if not pattern:
                                continue
                            name = str(item.get("name") or item.get("label") or "custom")
                            sanitized.append({"pattern": pattern, "name": name})
                    if sanitized:
                        safe_kwargs[key] = sanitized
                else:
                    safe_kwargs[key] = value

            logger.info(
                "Обрабатываем документ %s операцией %s для пользователя %s",
                original_name,
                operation,
                user_id,
            )
            if safe_kwargs:
                logger.debug("Применяем параметры операции: %s", safe_kwargs)

            result = await processor.safe_process(document_info.file_path, **safe_kwargs)

            if result.success:
                result.data["document_info"] = {
                    "original_name": document_info.original_name,
                    "file_size": document_info.size,
                    "upload_time": document_info.upload_time.isoformat(),
                    "user_id": user_id,
                    "local_path": str(document_info.file_path),
                }
                if document_info.remote_path:
                    result.data["document_info"]["remote_path"] = document_info.remote_path
                if safe_kwargs:
                    result.data["applied_options"] = dict(safe_kwargs)

                exports = self._create_exports(operation, result.data, document_info, safe_kwargs)
                if exports:
                    result.data["exports"] = exports

            await self.storage.cleanup_old_files(user_id, max_age_hours=24)
            return result

        except Exception as exc:
            logger.error("Ошибка обработки документа %s: %s", original_name, exc)
            raise ProcessingError(f"Ошибка обработки документа: {exc}", "PROCESSING_ERROR")

    def _get_processor(self, operation: str):
        processors = {
            "summarize": self.summarizer,
            "analyze_risks": self.risk_analyzer,
            "chat": self.document_chat,
            "anonymize": self.anonymizer,
            "translate": self.translator,
            "ocr": self.ocr_converter,
        }
        return processors.get(operation)

    def _append_export_note(self, base_text: str, data: dict[str, Any]) -> str:
        exports = data.get("exports") or []
        if not exports:
            return base_text

        lines = []
        for export in exports:
            fmt = str(export.get("format", "file")).upper()
            path_value = export.get("path", "")
            file_name = Path(path_value).name if path_value else path_value
            lines.append(f"• {html_escape(fmt)}: {html_escape(file_name)}")

        return base_text + "\n\n📎 Доступные файлы:\n" + "\n".join(lines)

    def _create_exports(
        self,
        operation: str,
        data: dict[str, Any],
        document_info,
        options: dict[str, Any],
    ) -> list[dict[str, Any]]:
        exports: list[dict[str, Any]] = []
        try:
            if operation == "summarize":
                formats = options.get("output_formats") or ["docx"]
                exports.extend(self._export_summary(document_info, data, formats))
            elif operation == "translate":
                formats = options.get("output_formats") or ["docx", "txt"]
                exports.extend(self._export_translation(document_info, data, formats))
            elif operation == "ocr":
                output_format = options.get("output_format", "txt")
                exports.extend(self._export_ocr(document_info, data, output_format))
            elif operation == "analyze_risks":
                exports.extend(self._export_risk_report(document_info, data))
            elif operation == "anonymize":
                exports.extend(self._export_anonymized(document_info, data))
        except Exception as export_error:
            logger.warning("Не удалось подготовить экспорт для %s: %s", operation, export_error)
        return exports

    def _export_summary(
        self, document_info, data: dict[str, Any], formats: list[str]
    ) -> list[dict[str, Any]]:
        summary_content = data.get("summary", {}).get("content")
        if not summary_content:
            return []

        exports: list[dict[str, Any]] = []
        base_name = f"{document_info.file_path.stem}_summary"
        export_dir = document_info.file_path.parent

        if "docx" in formats:
            if not self._dependency_available('docx'):
                exports.append(
                    self._dependency_notice(
                        dependency='python-docx',
                        feature='DOCX export',
                        format_name='docx',
                    )
                )
            else:
                from docx import Document  # type: ignore

                doc = Document()
                doc.add_heading("Сводка по документу", level=1)
                for block in summary_content.split("\n\n"):
                    doc.add_paragraph(block)
                docx_path = export_dir / f"{base_name}.docx"
                doc.save(docx_path)
                exports.append({"path": str(docx_path), "format": "docx", "label": "Резюме (DOCX)"})

        return exports

    def _export_anonymized(
        self, document_info, data: dict[str, Any]
    ) -> list[dict[str, Any]]:
        anonymized_text = data.get("anonymized_text")
        if not anonymized_text:
            return []

        exports: list[dict[str, Any]] = []
        base_name = f"{document_info.file_path.stem}_anonymized"
        export_dir = document_info.file_path.parent

        if not self._dependency_available('docx'):
            exports.append(
                self._dependency_notice(
                    dependency='python-docx',
                    feature='Anonymized DOCX export',
                    format_name='docx',
                )
            )
            return exports

        try:
            docx_path = export_dir / f"{base_name}.docx"
            original_path = Path(document_info.file_path)
            if not self._render_text_with_original_formatting(original_path, anonymized_text, docx_path):
                self._create_plain_docx(anonymized_text, docx_path)
            exports.append({"path": str(docx_path), "format": "docx", "label": "Обезличенный DOCX"})
        except Exception as docx_error:
            logger.warning("Failed to build anonymized DOCX: %s", docx_error)
            exports.append({"format": "docx", "error": "Не удалось сформировать DOCX экспорт"})

        return exports

    def _export_translation(
        self, document_info, data: dict[str, Any], formats: list[str]
    ) -> list[dict[str, Any]]:
        translated_text = data.get("translated_text") or ""
        if not translated_text:
            return []

        exports: list[dict[str, Any]] = []
        base_name = f"{document_info.file_path.stem}_translation"
        export_dir = document_info.file_path.parent
        original_path = Path(document_info.file_path)

        if "docx" in formats:
            if not self._dependency_available('docx'):
                exports.append(
                    self._dependency_notice(
                        dependency='python-docx',
                        feature='Translation DOCX export',
                        format_name='docx',
                    )
                )
            else:
                try:
                    docx_path = export_dir / f"{base_name}.docx"
                    self._create_translated_docx(original_path, translated_text, docx_path)
                    exports.append({"path": str(docx_path), "format": "docx", "label": "Перевод (DOCX)"})
                except Exception as exc:
                    logger.error("Не удалось сохранить перевод с форматированием: %s", exc)
                    fallback_path = export_dir / f"{base_name}_plain.docx"
                    self._create_plain_docx(translated_text, fallback_path)
                    exports.append({"path": str(fallback_path), "format": "docx", "label": "Перевод (DOCX)"})

        if "txt" in formats:
            txt_path = export_dir / f"{base_name}.txt"
            self._write_text_file(txt_path, translated_text)
            exports.append({"path": str(txt_path), "format": "txt", "label": "Перевод (TXT)"})

        return exports

    def _export_ocr(
        self, document_info, data: dict[str, Any], output_format: str
    ) -> list[dict[str, Any]]:
        recognized_text = data.get("recognized_text")
        if not recognized_text:
            return []

        exports: list[dict[str, Any]] = []
        base_name = f"{document_info.file_path.stem}_ocr"
        export_dir = document_info.file_path.parent
        fmt = output_format.lower()

        if fmt == "txt":
            txt_path = export_dir / f"{base_name}.txt"
            self._write_text_file(txt_path, recognized_text)
            exports.append({"path": str(txt_path), "format": "txt", "label": "Оцифровка (TXT)"})
        elif fmt == "docx":
            if not self._dependency_available('docx'):
                exports.append(
                    self._dependency_notice(
                        dependency='python-docx',
                        feature='OCR DOCX export',
                        format_name='docx',
                    )
                )
            else:
                from docx import Document  # type: ignore

                doc = Document()
                for block in recognized_text.split("\n\n"):
                    doc.add_paragraph(block)
                docx_path = export_dir / f"{base_name}.docx"
                doc.save(docx_path)
                exports.append({"path": str(docx_path), "format": "docx", "label": "Оцифровка (DOCX)"})

        elif fmt == "pdf":
            if not self._dependency_available('reportlab'):
                exports.append(
                    self._dependency_notice(
                        dependency='reportlab',
                        feature='OCR PDF export',
                        format_name='pdf',
                    )
                )
            else:
                from reportlab.lib.pagesizes import A4  # type: ignore
                from reportlab.lib.units import mm  # type: ignore
                from reportlab.pdfgen import canvas  # type: ignore

                pdf_path = export_dir / f"{base_name}.pdf"
                canv = canvas.Canvas(str(pdf_path), pagesize=A4)
                width, height = A4
                text_obj = canv.beginText(20 * mm, height - 20 * mm)
                for line in recognized_text.splitlines():
                    text_obj.textLine(line)
                    if text_obj.getY() < 20 * mm:
                        canv.drawText(text_obj)
                        canv.showPage()
                        text_obj = canv.beginText(20 * mm, height - 20 * mm)
                canv.drawText(text_obj)
                canv.save()
                exports.append({"path": str(pdf_path), "format": "pdf", "label": "Оцифровка (PDF)"})


        else:
            logger.warning("Неподдерживаемый формат экспорта OCR: %s", output_format)

        return exports


    def _export_risk_report(self, document_info, data: dict[str, Any]) -> list[dict[str, Any]]:
        overall = data.get("overall_risk_level")
        pattern_risks = data.get("pattern_risks", []) or []
        ai_section = data.get("ai_analysis", {}) or {}
        ai_risks = ai_section.get("risks", []) or []
        summary = ai_section.get("summary") or ""
        highlighted = data.get("highlighted_text", "") or ""
        recommendations = data.get("recommendations", []) or []
        compliance = data.get("legal_compliance", {}) or {}
        violations = compliance.get("violations") or []
        timestamp = data.get("analysis_timestamp")

        if not (
            overall
            or pattern_risks
            or ai_risks
            or violations
            or highlighted
            or summary
            or recommendations
        ):
            return []

        exports: list[dict[str, str]] = []
        base_name = f"{document_info.file_path.stem}_risk_report"
        export_dir = document_info.file_path.parent

        docx_created = False
        if self._dependency_available('docx'):
            try:
                from docx import Document  # type: ignore

                doc = Document()
                doc.add_heading('Отчёт о рисках', level=1)
                doc.add_paragraph(f"Документ: {document_info.original_name}")
                if timestamp:
                    try:
                        dt = datetime.fromisoformat(str(timestamp))
                        doc.add_paragraph(f"Дата анализа: {dt.strftime('%d.%m.%Y %H:%M')}")
                    except (ValueError, TypeError):
                        doc.add_paragraph(f"Дата анализа: {timestamp}")
                if overall:
                    level_text = self._localize_risk_level(overall, capitalize=True)
                    doc.add_paragraph(f"Общий уровень риска: {level_text}")
                if summary:
                    doc.add_heading('Краткое резюме', level=2)
                    doc.add_paragraph(summary)

                combined_risks = pattern_risks + ai_risks
                if combined_risks:
                    doc.add_heading('Обнаруженные риски', level=2)
                    for risk in combined_risks[:30]:
                        level_text = self._localize_risk_level(risk.get('risk_level'), capitalize=True)
                        description = str(risk.get('description') or 'Описание не указано')
                        paragraph = doc.add_paragraph(style='List Bullet')
                        run = paragraph.add_run(f"{level_text}: ")
                        run.bold = True
                        paragraph.add_run(description)
                        clause = risk.get('clause_text')
                        if clause:
                            doc.add_paragraph(str(clause))
                        refs = [str(ref) for ref in (risk.get('law_refs') or []) if ref]
                        if refs:
                            doc.add_paragraph('Нормативные ссылки: ' + ', '.join(refs))
                    if len(combined_risks) > 30:
                        doc.add_paragraph(f"...и ещё {len(combined_risks) - 30} рисков", style='List Bullet')

                if violations:
                    doc.add_heading('Потенциальные нарушения', level=2)
                    for violation in violations[:30]:
                        paragraph = doc.add_paragraph(style='List Bullet')
                        note = str(violation.get('note') or violation.get('text') or 'Нарушение')
                        paragraph.add_run(note)
                        refs = [str(ref) for ref in (violation.get('law_refs') or []) if ref]
                        if refs:
                            doc.add_paragraph('Нормативные ссылки: ' + ', '.join(refs))
                    if len(violations) > 30:
                        doc.add_paragraph(f"...и ещё {len(violations) - 30} пунктов", style='List Bullet')

                if recommendations:
                    doc.add_heading('Рекомендации', level=2)
                    for rec in recommendations[:20]:
                        doc.add_paragraph(str(rec), style='List Bullet')

                if highlighted:
                    doc.add_heading('Контекст с подсветкой', level=2)
                    doc.add_paragraph(str(highlighted))

                docx_path = export_dir / f"{base_name}.docx"
                doc.save(docx_path)
                exports.append({"path": str(docx_path), "format": "docx", "label": "Отчёт о рисках (DOCX)"})
                docx_created = True
            except Exception as exc:  # pragma: no cover
                logger.error("Не удалось сформировать DOCX отчёт о рисках: %s", exc)

        if not docx_created:
            txt_path = export_dir / f"{base_name}.txt"
            report_lines = [f"Общий уровень риска: {overall or 'неизвестен'}", "", "Найденные риски:"]
            for risk in pattern_risks:
                report_lines.append(
                    f"- [{risk.get('risk_level', 'unknown')}] {risk.get('description', '')}"
                )
            for risk in ai_risks:
                report_lines.append(
                    f"- [{risk.get('risk_level', 'unknown')}] {risk.get('description', '')}"
                )
            if summary:
                report_lines.extend(["", "Краткое резюме:", summary])
            if recommendations:
                report_lines.extend(["", "Рекомендации:"] + [f"- {rec}" for rec in recommendations])
            if highlighted:
                report_lines.extend(["", "Контекст с подсветкой:", highlighted])
            self._write_text_file(txt_path, "\n".join(report_lines))
            exports.append({"path": str(txt_path), "format": "txt", "label": "Отчёт о рисках"})

        return exports

    def _write_text_file(self, path: Path, content: str) -> None:
        path.write_text(content or "", encoding="utf-8")

    async def chat_with_document(
        self, user_id: int, document_id: str, question: str
    ) -> dict[str, Any]:
        try:
            return await self.document_chat.chat_with_document(document_id, question)
        except Exception as exc:
            logger.error("Ошибка чата с документом %s: %s", document_id, exc)
            raise ProcessingError(f"Ошибка чата: {exc}", "CHAT_ERROR")

    def get_supported_operations(self) -> dict[str, dict[str, Any]]:
        return {
            "summarize": {
                "name": "Саммаризация",
                "description": "Создание краткой выжимки документа с ключевыми положениями",
                "emoji": "📋",
                "formats": [".pdf", ".docx", ".doc", ".txt"],
                "parameters": ["detail_level", "language"],
            },
            "analyze_risks": {
                "name": "Анализ рисков",
                "description": "Выявление рисков и проблем в договорах и документах",
                "emoji": "⚠️",
                "formats": [".pdf", ".docx", ".doc", ".txt"],
                "parameters": ["custom_criteria"],
            },
            "chat": {
                "name": "Чат с документом",
                "description": "Интерактивные вопросы-ответы по документу",
                "emoji": "💬",
                "formats": [".pdf", ".docx", ".doc", ".txt"],
                "parameters": [],
            },
            "anonymize": {
                "name": "Обезличивание",
                "description": "Удаление персональных данных из документа",
                "emoji": "🔐",
                "formats": [".pdf", ".docx", ".doc", ".txt"],
                "parameters": ["anonymization_mode", "exclude_types", "custom_patterns"],
            },
            "translate": {
                "name": "Перевод документов",
                "description": "Перевод документа на другие языки",
                "emoji": "🌍",
                "formats": [".pdf", ".docx", ".doc", ".txt"],
                "parameters": ["source_lang", "target_lang"],
            },
            "ocr": {
                "name": "OCR распознавание",
                "description": "Распознавание текста из сканированных документов",
                "emoji": "🖭",
                "formats": [".pdf", ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"],
                "parameters": ["output_format"],
            },
        }

    def get_operation_info(self, operation: str) -> dict[str, Any] | None:
        return self.get_supported_operations().get(operation)

    def build_telegram_chunks(self, html: str, max_len: int = 3900) -> list[str]:
        """Собирает безопасный HTML и режет на куски под лимиты Telegram."""
        safe_html = format_safe_html(html)               # нормализуем/балансируем теги
        chunks = split_html_for_telegram(safe_html, max_len=max_len)
        return chunks

    def format_result_for_telegram(self, result: DocumentResult, operation: str) -> str:
        if not result.success:
            raw_html = f"✖ <b>Ошибка обработки</b>\n\n{html_escape(str(result.message))}"
            return format_safe_html(raw_html)

        operation_info = self.get_operation_info(operation) or {}
        emoji = operation_info.get("emoji", "📄")
        name = operation_info.get("name", operation.title())

        header = f"{emoji} <b>{html_escape(name)}</b>\n"
        header += f"⏱️ Время обработки: {result.processing_time:.1f}с\n\n"

        # Получаем сырой HTML от соответствующего форматтера
        if operation == "summarize":
            raw_html = self._format_summary_result(header, result.data)
        elif operation == "analyze_risks":
            raw_html = self._format_risk_analysis_result(header, result.data)
        elif operation == "chat":
            raw_html = self._format_chat_result(header, result.data)
        elif operation == "anonymize":
            raw_html = self._format_anonymize_result(header, result.data)
        elif operation == "translate":
            raw_html = self._format_translate_result(header, result.data)
        elif operation == "ocr":
            raw_html = self._format_ocr_result(header, result.data)
        else:
            raw_html = f"{header}✔ {result.message}"

        # Конвертируем переводы строк и очищаем HTML перед отправкой
        normalized = raw_html.replace("\r\n", "\n")
        sanitized = sanitize_telegram_html(normalized)
        sanitized = sanitized.replace("<br><br>", "\n\n").replace("<br>", "\n")
        return sanitized

    def _format_summary_result(self, header: str, data: dict[str, Any]) -> str:
        summary = data.get("summary", {})
        content = summary.get("content", "Содержимое недоступно")
        metadata = data.get("metadata", {})
        detail_level = data.get("detail_level", "detailed")
        language = data.get("language", "ru")
        detail_text = {"detailed": "Подробная", "brief": "Краткая"}.get(detail_level, detail_level)
        language_text = {"ru": "Русский", "en": "Английский"}.get(language, language)

        result = f"{header}<b>📝 Саммаризация:</b>\n"
        result += f"Уровень детализации: {html_escape(detail_text)}\n"
        result += f"Язык: {html_escape(language_text)}\n\n{html_escape(content)}"

        if metadata:
            result += "\n\n<b>📊 Метаданные:</b>\n"
            for key, value in metadata.items():
                result += f"• {html_escape(str(key))}: {html_escape(str(value))}\n"

        return self._append_export_note(result, data)


    def _format_risk_analysis_result(self, header: str, data: dict[str, Any]) -> str:
        overall_risk = data.get("overall_risk_level")
        risk_emojis = {"low": "🟢", "medium": "🟡", "high": "🟠", "critical": "🔴"}
        level_key = str(overall_risk).lower() if overall_risk is not None else ""
        emoji = risk_emojis.get(level_key, "✅")
        localized_level = self._localize_risk_level(overall_risk, uppercase=True)
        result = f"{header}<b>{emoji} Общий уровень риска: {html_escape(localized_level)}</b>\n\n"

        pattern_risks = data.get("pattern_risks", []) or []
        ai_risks = data.get("ai_analysis", {}).get("risks", []) or []
        combined_risks = pattern_risks + ai_risks
        if combined_risks:
            result += f"<b>⚠️ Обнаруженные риски: {len(combined_risks)}</b>\n\n"
            for risk in combined_risks[:3]:
                level_emoji = risk_emojis.get(str(risk.get('risk_level', 'medium')).lower(), "🟡")
                level_label = html_escape(self._localize_risk_level(risk.get('risk_level'), capitalize=True))
                desc = html_escape(str(risk.get('description', 'Неописанный риск')))
                result += f"{level_emoji} <b>{level_label}</b>: {desc}\n"
            if len(combined_risks) > 3:
                result += f"\n<i>...и ещё {len(combined_risks) - 3} рисков</i>\n"
        else:
            result += "Риски не обнаружены\n"

        ai_analysis = data.get("ai_analysis", {}) or {}
        summary = ai_analysis.get("summary")
        if summary:
            trimmed = summary if len(summary) <= 1500 else summary[:1500] + "...\n\n(Полный отчёт доступен в файле)"
            result += f"\n<b>🤖 Краткое резюме:</b>\n{html_escape(trimmed)}\n"

        legal_compliance = data.get("legal_compliance", {}) or {}
        violations = legal_compliance.get("violations") or []
        if violations:
            result += "\n<b>⚖️ Возможные нарушения:</b>\n"
            for violation in violations[:3]:
                text_value = str(violation.get("text") or violation.get("note") or "Нарушение")
                text_line = html_escape(text_value)
                refs = [str(ref) for ref in (violation.get("law_refs") or []) if ref]
                if refs:
                    refs_text = html_escape(", ".join(refs))
                    text_line += f" ({refs_text})"
                result += f"- {text_line}\n"
            if len(violations) > 3:
                result += f"<i>...и ещё {len(violations) - 3} пунктов</i>\n"
        elif legal_compliance.get("status") == "completed":
            result += "\n<b>⚖️ Возможные нарушения:</b> не выявлены.\n"

        recommendations = data.get("recommendations", []) or []
        if recommendations:
            result += "\n<b>Рекомендации:</b>\n"
            for rec in recommendations[:3]:
                result += f"• {html_escape(str(rec))}\n"

        return self._append_export_note(result, data)

    def _format_chat_result(self, header: str, data: dict[str, Any]) -> str:
        question = data.get("question", "")
        answer = data.get("answer", "Ответ недоступен")

        result = f"{header}<b>✔ Вопрос:</b> {html_escape(question)}\n\n"
        result += f"<b>💡 Ответ:</b>\n{html_escape(answer)}\n\n"

        context_chunks = data.get("context_chunks", [])
        if context_chunks:
            result += "<b>🔍 Использованные фрагменты:</b>\n"
            for item in context_chunks:
                preview = item.get("excerpt", "")
                if len(preview) > 160:
                    preview = preview[:160] + "..."
                score = item.get("score")
                score_text = f" (релевантность {score:.2f})" if isinstance(score, (int, float)) else ""
                result += f"• Фрагмент {int(item.get('index', 0)) + 1}{score_text}: <i>{html_escape(preview)}</i>\n"
            result += "\n"

        fragments = data.get("relevant_fragments", [])
        if fragments:
            result += "<b>📓 Релевантные предложения:</b>\n"
            for i, fragment in enumerate(fragments[:2], 1):
                text = html_escape(fragment.get('text', '')[:200])
                result += f"{i}. <i>{text}...</i>\n"

        return self._append_export_note(result, data)

    def _format_anonymize_result(self, header: str, data: dict[str, Any]) -> str:
        report = data.get("anonymization_report", {})
        stats = report.get("statistics", {})
        result = f"{header}<b>🔒 Результаты обезличивания:</b>\n\n"

        total_items = sum(int(v) for v in stats.values()) if stats else 0
        result += f"Всего обработано элементов: <b>{total_items}</b>\n\n"

        if stats:
            result += "<b>По типам данных:</b>\n"
            type_names = {
                "names": "👤 ФИО",
                "phones": "📞 Телефоны",
                "emails": "📧 Email",
                "addresses": "🏠 Адреса",
                "documents": "📄 Номера документов",
                "bank_details": "🏦 Банковские реквизиты",
                "badge_numbers": "🆔 Табельные номера",
                "registration_numbers": "🗂️ Регистрационные записи",
                "domains": "🌐 Домены",
                "urls": "🔗 Ссылки",
            }
            label_map = report.get("type_labels", {})
            for data_type, count in stats.items():
                if int(count) > 0:
                    name = label_map.get(data_type) or type_names.get(data_type) or data_type
                    result += f"• {html_escape(str(name))}: {int(count)}\n"

        custom_used = data.get("applied_custom_patterns") or []
        if custom_used:
            result += "\n<b>Пользовательские шаблоны:</b>\n"
            for entry in custom_used:
                label = entry.get("label") or entry.get("kind")
                result += f"• {html_escape(str(label))}\n"

        result += "\n✅ Документ готов к безопасной передаче"
        return self._append_export_note(result, data)

    def _format_translate_result(self, header: str, data: dict[str, Any]) -> str:
        source_lang = data.get("source_language", "")
        target_lang = data.get("target_language", "")
        lang_names = {
            "ru": "🇷🇺 Русский",
            "en": "🇺🇸 Английский",
            "zh": "🇨🇳 Китайский",
            "de": "🇩🇪 Немецкий",
        }
        source_name = lang_names.get(source_lang, source_lang)
        target_name = lang_names.get(target_lang, target_lang)

        metadata = data.get("translation_metadata", {})

        result = f"{header}<b>🌍 Перевод завершен</b>\n"
        result += f"{html_escape(source_name)} → {html_escape(target_name)}\n\n"

        if metadata:
            result += "<b>📊 Статистика:</b>\n"
            result += f"• Исходный текст: {int(metadata.get('original_length', 0))} символов\n"
            result += f"• Переведенный текст: {int(metadata.get('translated_length', 0))} символов\n"
            if metadata.get('chunks_processed'):
                result += f"• Частей обработано: {int(metadata['chunks_processed'])}\n"
            result += "\n"

        translated_text = data.get("translated_text", "") or ""
        if len(translated_text) > 2000:
            preview = translated_text[:2000] + "..."
            result += f"<b>📄 Предварительный просмотр:</b>\n{html_escape(preview)}\n\n(Полный перевод доступен в файле)"
        else:
            result += f"<b>📄 Переведенный текст:</b>\n{html_escape(translated_text)}"

        return self._append_export_note(result, data)

    def _format_ocr_result(self, header: str, data: dict[str, Any]) -> str:
        confidence = data.get("confidence_score", 0)
        quality = data.get("quality_analysis", {})

        confidence_emoji = "🟢" if confidence >= 80 else "🟡" if confidence >= 60 else "🔴"

        result = f"{header}<b>👁️ OCR распознавание завершено</b>\n"
        result += f"{confidence_emoji} Уверенность: {confidence:.1f}%\n"
        result += f"📊 Качество: {html_escape(quality.get('quality_level', 'неизвестно'))}\n\n"

        processing_info = data.get("processing_info", {})
        if processing_info:
            result += "<b>📈 Статистика:</b>\n"
            result += f"• Слов распознано: {processing_info.get('word_count', 0)}\n"
            result += f"• Символов: {processing_info.get('text_length', 0)}\n\n"

        recognized_text = data.get("recognized_text", "")
        if len(recognized_text) > 2000:
            preview = html_escape(recognized_text[:2000]) + "..."
            result += f"<b>📄 Распознанный текст:</b>\n{preview}\n\n<i>(Полный текст доступен в файле)</i>"
        else:
            result += f"<b>📄 Распознанный текст:</b>\n{html_escape(recognized_text)}"

        recommendations = quality.get("recommendations", [])
        if recommendations:
            result += "\n\n<b>💡 Рекомендации:</b>\n"
            for rec in recommendations[:2]:
                result += f"• {html_escape(rec)}\n"

        return self._append_export_note(result, data)

    async def cleanup_user_files(self, user_id: int, max_age_hours: int = 24):
        await self.storage.cleanup_old_files(user_id, max_age_hours)
        await asyncio.to_thread(self.document_chat.cleanup_old_documents, max_age_hours)

    def _create_translated_docx(self, original_path: Path, translated_text: str, export_path: Path) -> None:
        if not self._render_text_with_original_formatting(original_path, translated_text, export_path):
            self._create_plain_docx(translated_text, export_path)

    def _render_text_with_original_formatting(self, original_path: Path, text: str, export_path: Path) -> bool:
        if original_path.suffix.lower() != ".docx" or not original_path.exists():
            return False

        from docx import Document  # type: ignore

        doc = Document(original_path)
        paragraphs = self._collect_paragraphs(doc)
        if not paragraphs:
            return False

        segments = self._split_text_segments(len(paragraphs), text)
        for paragraph, segment in zip(paragraphs, segments):
            self._replace_paragraph_text(paragraph, segment)
        doc.save(export_path)
        return True

    def _collect_paragraphs(self, document) -> List[Any]:
        paragraphs: List[Any] = list(document.paragraphs)
        for table in document.tables:
            paragraphs.extend(self._collect_table_paragraphs(table))
        return paragraphs

    def _collect_table_paragraphs(self, table) -> List[Any]:
        paragraphs: List[Any] = []
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
                for inner_table in cell.tables:
                    paragraphs.extend(self._collect_table_paragraphs(inner_table))
        return paragraphs

    def _split_text_segments(self, paragraph_count: int, text: str) -> List[str]:
        if paragraph_count == 0:
            raise ValueError("Document has no paragraphs")

        normalized = (text or "").replace("\r\n", "\n")

        double_segments = [seg.strip("\n") for seg in normalized.split("\n\n")]
        if len(double_segments) == paragraph_count:
            return double_segments
        if len(double_segments) == paragraph_count + 1 and double_segments[-1] == "":
            return double_segments[:-1]

        single_segments = [seg.strip("\n") for seg in normalized.split("\n")]
        if len(single_segments) == paragraph_count:
            return single_segments
        if len(single_segments) == paragraph_count + 1 and single_segments[-1] == "":
            return single_segments[:-1]

        if len(single_segments) < paragraph_count:
            single_segments.extend([""] * (paragraph_count - len(single_segments)))
            return single_segments[:paragraph_count]
        if paragraph_count == 1:
            return ["\n".join(single_segments)]

        head = single_segments[: paragraph_count - 1]
        tail = "\n".join(single_segments[paragraph_count - 1 :])
        head.append(tail)
        return head

    def _replace_paragraph_text(self, paragraph, text: str) -> None:
        style = paragraph.style
        for run in list(paragraph.runs):
            paragraph._element.remove(run._element)
        paragraph.style = style

        lines = text.split("\n") if text else [""]
        for index, line in enumerate(lines):
            run = paragraph.add_run(line)
            if index < len(lines) - 1:
                run.add_break()

    def _create_plain_docx(self, text: str, export_path: Path) -> None:
        from docx import Document  # type: ignore

        doc = Document()
        normalized = (text or "").replace("\r\n", "\n")
        blocks = normalized.split("\n\n")
        for block in blocks:
            doc.add_paragraph(block)
        doc.save(export_path)
