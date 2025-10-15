"""
Модуль обезличивания (анонимизации) документов
Удаление персональных данных из документов для безопасного обмена.

Особенности:
- Один проход по тексту с единым комбинированным регэкспом (без «двойных замен»).
- Поддержка: ФИО, телефоны (междунар.), email (в т.ч. unicode \w), адреса (RU-маркеры),
  документы РФ (паспорт 4+6, СНИЛС с чексуммой, ИНН 10/12 с чексуммой), банковские
  реквизиты (р/с, БИК), карты (Luhn), IBAN (mod-97), даты рождения dd.mm.yyyy.
- Режимы: replace / mask (сохранение формата для телефонов/email) / remove / pseudonym.
- Псевдонимы детерминированы (HMAC-SHA256) и стабильны между запусками при ANON_SECRET.
- Отчёт: статистика по типам, список обработанных фрагментов со сниппетами и координатами.
"""

from __future__ import annotations

import base64 as _b64
import json
import hmac
import hashlib as _hash
import logging
from src.core.settings import AppSettings

import re
from dataclasses import dataclass
from pathlib import Path
from textwrap import dedent
from typing import Any, Awaitable, Callable, Dict, List, Tuple, TYPE_CHECKING

from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from .base import DocumentProcessor, DocumentResult, ProcessingError
from .utils import FileFormatHandler, TextProcessor

if TYPE_CHECKING:  # pragma: no cover - hints
    from src.core.openai_service import OpenAIService

logger = logging.getLogger(__name__)


# ---------------------------- ВАЛИДАЦИИ / УТИЛИТЫ ----------------------------

def _digits(s: str) -> str:
    return "".join(ch for ch in s if ch.isdigit())


def _luhn_ok(digits_only: str) -> bool:
    """Проверка Luhn для банковских карт (13–19 цифр)."""
    if not digits_only or not digits_only.isdigit():
        return False
    total = 0
    rev = digits_only[::-1]
    for i, ch in enumerate(rev):
        d = ord(ch) - 48
        if i % 2 == 1:
            d *= 2
            if d > 9:
                d -= 9
        total += d
    return total % 10 == 0


def _snils_ok(snils: str) -> bool:
    """СНИЛС: XXX-XXX-XXX YY (YY — контрольная сумма)."""
    d = _digits(snils)
    if len(d) != 11:
        return False
    num, cs = d[:9], int(d[9:])
    s = sum((9 - i) * int(num[i]) for i in range(9))
    if s < 100:
        chk = s
    elif s in (100, 101):
        chk = 0
    else:
        chk = s % 101
        if chk == 100:
            chk = 0
    return chk == cs


def _inn_ok(inn: str) -> bool:
    """ИНН РФ: 10 (юр) или 12 (физ), с контрольными цифрами."""
    d = _digits(inn)
    if len(d) == 10:
        k = [2, 4, 10, 3, 5, 9, 4, 6, 8]
        chk = sum(int(d[i]) * k[i] for i in range(9)) % 11 % 10
        return chk == int(d[9])
    if len(d) == 12:
        k1 = [7, 2, 4, 10, 3, 5, 9, 4, 6, 8, 0]
        k2 = [3, 7, 2, 4, 10, 3, 5, 9, 4, 6, 8, 0]
        c1 = sum(int(d[i]) * k1[i] for i in range(11)) % 11 % 10
        c2 = sum(int(d[i]) * k2[i] for i in range(12)) % 11 % 10
        return c1 == int(d[10]) and c2 == int(d[11])
    return False

def _is_passport(doc: str) -> bool:
    digits = _digits(doc)
    return len(digits) == 10




def _iban_ok(iban: str) -> bool:
    """IBAN mod-97 (упрощённая проверка формата + контроль)."""
    s = re.sub(r"\s+", "", iban).upper()
    if not re.fullmatch(r"[A-Z]{2}\\d{2}[A-Z0-9]{11,30}", s):
        return False
    s = s[4:] + s[:4]
    num = []
    for ch in s:
        if ch.isdigit():
            num.append(ch)
        else:
            num.append(str(ord(ch) - 55))
    n = "".join(num)
    rem = 0
    for i in range(0, len(n), 9):
        rem = int(str(rem) + n[i:i + 9]) % 97
    return rem == 1


def _mask_preserve(s: str) -> str:
    """Маска, сохраняющая разделители (буквы/цифры → *)."""
    return "".join("*" if ch.isalnum() else ch for ch in s)


def _mask_phone(s: str) -> str:
    """Маска телефона: сохраняем формат и последние 4 цифры."""
    digits = _digits(s)
    if not digits:
        return _mask_preserve(s)
    # Сколько звёздочек до последних 4 цифр
    to_mask = max(0, len(digits) - 4)
    out = []
    masked = 0
    kept = 0
    for ch in s:
        if ch.isdigit():
            if masked < to_mask:
                out.append("*")
                masked += 1
            else:
                # сохраняем последние 4 цифры в исходном порядке
                out.append(digits[to_mask + kept])
                kept += 1
        else:
            out.append(ch)
    return "".join(out)


def _mask_email(s: str) -> str:
    """Маска email: скрываем локальную часть, домен сохраняем."""
    if "@" not in s:
        return _mask_preserve(s)
    local, domain = s.split("@", 1)
    if len(local) <= 2:
        mlocal = "*" * len(local)
    else:
        mlocal = local[0] + "*" * (len(local) - 2) + local[-1]
    return f"{mlocal}@{domain}"


def _pseudo_id(kind: str, original: str, secret: bytes) -> str:
    """Детерминированный короткий псевдоним по HMAC-SHA256 (6 символов base64url)."""
    digest = hmac.new(secret, f"{kind}:{original}".encode("utf-8"), _hash.sha256).digest()
    code = _b64.urlsafe_b64encode(digest[:4]).decode("ascii").rstrip("=")
    return code.upper()


# ------------------------------- ПАТТЕРНЫ -------------------------------

@dataclass
class PatternSpec:
    kind: str
    pattern: str
    validate: Callable[[str], bool] | None = None  # опциональная доп. проверка
    label: str | None = None


# ------------------------------ ОСНОВНОЙ КЛАСС ------------------------------

from textwrap import dedent

_AI_SYSTEM_PROMPT = dedent("""

    Ты – специализированный эксперт по полной анонимизации юридических документов. Твоя единственная цель – 
    удалить все персонально идентифицирующие данные (PII), сохраняя при этом АБСОЛЮТНО ВСЕ: юридическую значимость, 
    структуру, смысл и **полное форматирование** документа.

    ## КРИТИЧЕСКИ ВАЖНО: СОХРАНЕНИЕ ФОРМАТИРОВАНИЯ

    ### ОБЯЗАТЕЛЬНОЕ ПРАВИЛО ФОРМАТИРОВАНИЯ:
    **Каждый элемент форматирования исходного документа должен быть сохранён на 100% в анонимизированной версии.**

    ### ЭЛЕМЕНТЫ ФОРМАТИРОВАНИЯ, КОТОРЫЕ ДОЛЖНЫ БЫТЬ СОХРАНЕНЫ:

    #### 1. ТИПОГРАФИКА И СТИЛИ ТЕКСТА
    - **Жирный шрифт** (bold) – если слово/фраза были жирными, замена тоже должна быть жирной
      - Пример: **Иванов И.И.** → **[ФИО-1]**
    - *Курсив* (italic) – сохраняй курсивное начертание
      - Пример: *ООО "Компания"* → *[КОМПАНИЯ-1]*
    - <u>Подчёркивание</u> (underline) – если текст подчёркнут, замена тоже подчёркнута
      - Пример: <u>Петров П.П.</u> → <u>[ФИО-2]</u>
    - ~~Зачёркнутый текст~~ (strikethrough) – сохраняй зачёркивания
    - `Моноширинный шрифт` (monospace) – для кодов, номеров
    - Комбинации стилей: ***жирный курсив***, **<u>жирный подчёркнутый</u>** и т.д.

    #### 2. РАЗМЕР И ЦВЕТ ШРИФТА
    - Размер шрифта (если больше/меньше стандартного) – сохраняй
    - Цвет текста – если PII выделены цветом, замена должна иметь тот же цвет
    - Фон текста (highlight) – жёлтый, серый и другие выделения

    #### 3. ВЫРАВНИВАНИЕ И ОТСТУПЫ
    - Выравнивание текста: по левому краю, по центру, по правому краю, по ширине
    - Отступы слева/справа (indentation)
    - Красная строка (first-line indent)
    - Межстрочный интервал
    - Интервалы до/после абзаца

    #### 4. ЗАГОЛОВКИ И РАЗДЕЛЫ
    - Уровни заголовков (Heading 1, Heading 2, и т.д.) – если имя в заголовке, замена сохраняет стиль заголовка
      - Пример: **# ДОГОВОР с Ивановым И.И.** → **# ДОГОВОР с [ФИО-1]**
    - Нумерация разделов (1., 1.1., 1.1.1., и т.д.)
    - Стили заголовков документа

    #### 5. СПИСКИ И НУМЕРАЦИЯ
    - Маркированные списки (•, ○, ■, -, * и т.д.)
    - Нумерованные списки (1, 2, 3 / a, b, c / i, ii, iii / А, Б, В)
    - Вложенность списков (подпункты)
    - Отступы пунктов списка
    - Форматирование внутри пунктов списка

    #### 6. ТАБЛИЦЫ
    - Структура таблицы (количество строк и столбцов)
    - Ширина столбцов
    - Объединение ячеек (merged cells)
    - Границы таблицы (толщина, цвет, стиль линий)
    - Заливка ячеек (background color)
    - Выравнивание текста в ячейках
    - Форматирование текста внутри ячеек (жирный, курсив и т.д.)
    - Заголовки таблиц

    #### 7. СПЕЦИАЛЬНЫЕ ЭЛЕМЕНТЫ
    - Колонтитулы (headers/footers) – анонимизируй PII, сохраняя структуру
    - Номера страниц
    - Сноски и концевые сноски (footnotes/endnotes)
    - Гиперссылки – если содержат PII, замени, но сохрани форматирование ссылки
    - Перекрёстные ссылки внутри документа
    - Разрывы страниц и разделов
    - Водяные знаки (если не содержат PII)

    #### 8. СПЕЦИАЛЬНЫЕ СИМВОЛЫ И РАЗДЕЛИТЕЛИ
    - Горизонтальные линии (----, ════, ─────)
    - Специальные символы (§, №, ©, ®, ™, и т.д.)
    - Пробелы и табуляция
    - Неразрывные пробелы
    - Мягкие переносы

    #### 9. ДОКУМЕНТЫ С ОСОБЫМ ФОРМАТИРОВАНИЕМ
    - **Markdown**: сохраняй всю разметку (#, **, *, -, >, и т.д.)
    - **HTML**: сохраняй теги форматирования (<b>, <i>, <u>, <strong>, <em>, <span>, и т.д.)
    - **LaTeX**: сохраняй команды форматирования (\\textbf{}, \\textit{}, и т.д.)
    - **RTF/DOCX**: сохраняй все стили и форматирование
    - **PDF**: максимально сохраняй визуальное форматирование

    ## КАТЕГОРИИ ДАННЫХ ДЛЯ ОБЯЗАТЕЛЬНОЙ АНОНИМИЗАЦИИ

    ### 1. ПЕРСОНАЛЬНЫЕ ДАННЫЕ ФИЗИЧЕСКИХ ЛИЦ
    - ФИО (полные и сокращённые), инициалы, псевдонимы, прозвища
    - Даты рождения, возраст (если позволяет идентификацию)
    - Пол (если в сочетании с другими данными ведёт к идентификации)
    - Место рождения, гражданство, национальность
    - Семейное положение, сведения о родственниках (имена, степени родства)
    - Образование (если указаны конкретные учебные заведения с датами)

    ### 2. КОНТАКТНЫЕ ДАННЫЕ
    - Номера телефонов (мобильные, стационарные, факсы)
    - Адреса электронной почты
    - Аккаунты мессенджеров (Telegram, WhatsApp, Viber и т.д.)
    - Социальные сети (профили, ссылки, никнеймы)
    - Почтовые адреса (полные и частичные)
    - Фактические адреса (места жительства, работы, регистрации)
    - GPS-координаты, геолокационные данные

    ### 3. ИДЕНТИФИКАЦИОННЫЕ НОМЕРА И КОДЫ
    - Паспортные данные (серия, номер, кем и когда выдан)
    - Иные удостоверения личности (водительские права, военные билеты и т.д.)
    - СНИЛС, ИНН физических лиц
    - Номера медицинских полисов, карт пациента
    - Биометрические данные (отпечатки пальцев, сканы радужки, фотографии)

    ### 4. ДАННЫЕ ОРГАНИЗАЦИЙ И ИП
    - Полные и сокращённые наименования компаний (если они уникально идентифицируют лицо)
    - ИНН, КПП, ОГРН, ОГРНИП юридических лиц и ИП
    - Должности и ФИО руководителей, представителей
    - Торговые марки и бренды (если связаны с идентификацией конкретного лица)
    - Адреса регистрации и фактические адреса компаний

    ### 5. ФИНАНСОВЫЕ И ДОГОВОРНЫЕ ДАННЫЕ
    - Номера банковских счетов, карт
    - БИК, корреспондентские счета, SWIFT, IBAN
    - Номера договоров, контрактов, соглашений
    - Регистрационные номера судебных дел
    - Кадастровые номера (если позволяют идентифицировать владельца)
    - Номера транспортных средств, VIN-коды

    ### 6. МЕДИЦИНСКИЕ И БИОМЕТРИЧЕСКИЕ ДАННЫЕ
    - Диагнозы, истории болезней
    - Результаты анализов и обследований с идентификаторами
    - Генетическая информация
    - Данные о физических особенностях, способных идентифицировать лицо

    ### 7. ЦИФРОВЫЕ СЛЕДЫ
    - IP-адреса
    - MAC-адреса
    - Cookies, идентификаторы сессий
    - Логины, пароли, ключи доступа

    ## ПРАВИЛА ЗАМЕНЫ ДАННЫХ С СОХРАНЕНИЕМ ФОРМАТИРОВАНИЯ

    ### Принципы замены:

    1. **Консистентность**: одни и те же данные заменяй одинаково в пределах документа
       - Пример: "Иванов И.И." → всегда "[ФИО-1]", а "Петров П.П." → всегда "[ФИО-2]"

    2. **Типизация**: используй понятные метки замены в квадратных скобках:
       - `[ФИО-1]`, `[ФИО-2]` – для разных физических лиц
       - `[КОМПАНИЯ-1]`, `[КОМПАНИЯ-2]` – для организаций
       - `[АДРЕС-1]` – для почтовых/фактических адресов
       - `[EMAIL-1]` – для электронной почты
       - `[ТЕЛЕФОН-1]` – для телефонов
       - `[ПАСПОРТ-1]` – для паспортных данных
       - `[ИНН-1]` – для ИНН
       - `[СЧЁТ-1]` – для банковских счетов
       - `[ДОГОВОР-1]` – для номеров договоров
       - `[ДЕЛО-1]` – для номеров дел
       - `[ДАТА-1]` – для дат рождения и значимых персональных дат
       - `[ДОЛЖНОСТЬ-1]` – для должностей, привязанных к конкретному лицу

    3. **ФОРМАТИРОВАНИЕ ПЕРЕНОСИТСЯ**: если заменяемые данные имели форматирование, замена ОБЯЗАНА иметь то же форматирование
       - **Иванов И.И.** → **[ФИО-1]** (жирный)
       - *ООО "Ромашка"* → *[КОМПАНИЯ-1]* (курсив)
       - <u>+7 (999) 123-45-67</u> → <u>[ТЕЛЕФОН-1]</u> (подчёркнутый)
       - ***Директор Петров П.П.*** → ***Директор [ФИО-2]*** (жирный курсив)

    4. **Сохранение контекста**: если удаление данных сделает фразу бессмысленной, замени на общую категорию:
       - "Генеральный директор Иванов И.И." → "Генеральный директор [ФИО-1]"
       - "проживающий по адресу: г. Москва, ул. Ленина, д. 5" → "проживающий по адресу: [АДРЕС-1]"

    5. **Частичная анонимизация дат**: 
       - Даты рождения → всегда `[ДАТА-1]`
       - Даты документов, событий, не связанных с идентификацией → оставляй без изменений


  ## ОБЯЗАТЕЛЬНЫЕ ПРАВИЛА ОБРАБОТКИ

    ### ЧТО СОХРАНЯТЬ НА 100%:
    ✓ **Всю структуру документа** (разделы, главы, параграфы, пункты, подпункты)
    ✓ **Нумерацию** (1, 2, 3, а), б), в), i, ii, iii и т.д.)
    ✓ **Таблицы** (структура, границы, заливка, объединение ячеек)
    ✓ **Списки** (маркированные, нумерованные, вложенность)
    ✓ **Юридические термины**, формулировки, цитаты из законов
    ✓ **ВСЁ ФОРМАТИРОВАНИЕ**: жирный, курсив, подчёркивание, зачёркивание, цвет, размер, выравнивание
    ✓ **Пунктуацию**, регистр букв, язык оригинала
    ✓ **Статьи законов**, ссылки на НПА (нормативно-правовые акты)
    ✓ **Суммы, сроки, проценты**, не привязанные к конкретным лицам
    ✓ **Колонтитулы**, номера страниц, сноски
    ✓ **Специальные символы** (§, №, ©, ® и т.д.)
    ✓ **Разделители** (линии, рамки)
    ✓ **Отступы и интервалы**

    ### ЧТО ЗАПРЕЩЕНО:
    ✗ Сокращать или упрощать текст
    ✗ Делать резюме или выжимки
    ✗ Добавлять комментарии, пояснения, мета-информацию
    ✗ Изменять смысл юридических формулировок
    ✗ Добавлять вымышленные данные или домыслы
    ✗ Использовать техническую разметку (типа ~~удалено~~, [REDACTED] и т.п.)
    ✗ Менять язык документа
    ✗ **УДАЛЯТЬ ИЛИ ИЗМЕНЯТЬ ФОРМАТИРОВАНИЕ** – это критическая ошибка!
    ✗ Менять выравнивание текста
    ✗ Удалять пустые строки или изменять интервалы
    ✗ Упрощать таблицы или списки

    ### ОСОБЫЕ СЛУЧАИ:
    - **Если PII нет**: верни текст полностью без изменений, включая всё форматирование
    - **Если весь абзац – одни PII**: замени весь абзац на `[ПЕРСОНАЛЬНЫЕ_ДАННЫЕ_УДАЛЕНЫ]`, сохраняя форматирование абзаца (выравнивание, отступы)
    - **Косвенная идентификация**: если комбинация нейтральных данных позволяет идентифицировать лицо, анонимизируй эту комбинацию
    - **Подписи**: замени на `[ПОДПИСЬ]` или `[ПОДПИСЬ-ФИО-1]` для связи с конкретным лицом, сохраняя расположение подписи

    ## АЛГОРИТМ РАБОТЫ

    1. **Прочитай** документ целиком, обращая внимание на форматирование
    2. **Зафиксируй** все элементы форматирования для каждого PII
    3. **Выяви** все PII согласно категориям выше
    4. **Присвой** уникальные метки для каждого уникального элемента PII
    5. **Замени** все вхождения консистентно, **ПЕРЕНОСЯ ФОРМАТИРОВАНИЕ** на замену
    6. **Проверь**:
       - Не осталось ли незамеченных PII
       - Сохранён ли юридический смысл
       - Консистентны ли замены
       - **Сохранено ли всё форматирование на 100%**
       - Не нарушена ли структура таблиц и списков
       - Не изменились ли отступы и выравнивание
       - Не добавлен ли лишний текст

    ## ФОРМАТ ВЫДАЧИ

    Возвращай **ТОЛЬКО** анонимизированный текст документа в исходном формате со **100% сохранением форматирования**, без:
    - Вступительных фраз ("Вот анонимизированный документ...")
    - Заключительных комментариев ("Все PII удалены...")
    - Технических пометок
    - Пояснений о проделанной работе

    **АНОНИМИЗИРОВАННЫЙ ДОКУМЕНТ ДОЛЖЕН ВЫГЛЯДЕТЬ ВИЗУАЛЬНО ИДЕНТИЧНО ОРИГИНАЛУ, ОТЛИЧАЯСЬ ТОЛЬКО ЗАМЕНОЙ PII НА МЕТКИ.**

    ---

    **Важно**: Каждое нарушение конфиденциальности – это юридический риск. Каждое нарушение форматирования – это потеря юридической значимости документа. Твоя задача – абсолютная защита данных при сохранении полной юридической ценности и **полного визуального форматирования** документа.
    """
)


class DocumentAnonymizer(DocumentProcessor):
    """Класс для обезличивания персональных данных в документах"""

    def __init__(
        self,
        *,
        openai_service: "OpenAIService" | None = None,
        settings: AppSettings | None = None,
    ) -> None:
        super().__init__(name="DocumentAnonymizer", max_file_size=50 * 1024 * 1024)
        self.supported_formats = [".pdf", ".docx", ".doc", ".txt"]
        self.anonymization_map: dict[str, str] = {}

        self._openai_service = openai_service

        if settings is None:
            from src.core.app_context import get_settings  # avoid circular import

            settings = get_settings()
        self._settings = settings
        self._secret = settings.get_str("ANON_SECRET", "") or ""
        self._ai_enabled = bool(openai_service) and settings.get_bool("ANON_USE_AI", True)
        self._ai_chunk_chars = max(1000, settings.get_int("ANON_AI_CHUNK_CHARS", 3500))
        self._ai_chunk_overlap = max(0, settings.get_int("ANON_AI_CHUNK_OVERLAP", 200))
        self._ai_temperature = settings.get_float("ANON_AI_TEMPERATURE", 0.0)

        # Набор паттернов с минимизацией фолс-позитивов
        self._specs: List[PatternSpec] = [
            # ФИО: 2–3 слова, каждое ≥2 символов (отсечь «И.П.»)
            PatternSpec("names", r"\b[А-ЯЁ][а-яё]+(?:ов|ев|ёв|ин|ын|кий|цкий|ская|цкая|ова|ева|ёва|ина|ына|ский|ской)\b\s+[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+)?\b"),
            # Телефоны: международные/локальные, суммарно 10–15 цифр
            PatternSpec(
                "phones",
                r"(?:\+?\\d[\\d\-\s().]{6,}\\d)",
                validate=lambda s: 10 <= len(_digits(s)) <= 15,
            ),
            # Email: поддерживаем unicode \w в локальной части
            PatternSpec("emails", r"\b[\w.\-+%]+@[\w.\-]+\.[A-Za-zА-Яа-яЁё]{2,24}\b"),
            # Адреса: явные маркеры
            PatternSpec(
                "addresses",
                r"\b(?:г\.\s*[А-ЯЁ][а-яё\- ]+|ул\.\s*[А-ЯЁ][а-яё\- ]+|просп\.?\s*[А-ЯЁ][а-яё\- ]+|"
                r"пр-кт\.?\s*[А-ЯЁ][а-яё\- ]+|пер\.\s*[А-ЯЁ][а-яё\- ]+|дом\s*\\d+\w*|д\.\s*\\d+\w*)\b",
            ),
            # Документы РФ: паспорт 4+6, СНИЛС, ИНН 10/12 (с проверкой)
            PatternSpec(
                "documents",
                r"\b(?:(?:серия\s*)?\d{4}(?:\s*№\s*|\s+)?\d{6}|\d{3}-?\d{3}-?\d{3}\s?\d{2}|\d{10}|\d{12})\b",
                validate=lambda s: (
                    _snils_ok(s) or _inn_ok(s) or _is_passport(s)
                ),
            ),
            # Банковские реквизиты: р/с 20, БИК 9, карты 13–19 (Luhn)
            PatternSpec(
                "bank_details",
                r"(?:(?:р/с|к/с|р\s?с|к\s?с|бик|p/c|k/c)[:\s]*)?(?:\d[\d\s-]{11,}\d|\d{9}|\d{20})",
                validate=lambda s: (
                    len(_digits(s)) == 20
                    or len(_digits(s)) == 9
                    or (13 <= len(_digits(s)) <= 19 and _luhn_ok(_digits(s)))
                ),
            ),
            # IBAN (международные счета)
            PatternSpec("iban", r"\b[A-Z]{2}\\d{2}[A-Z0-9]{11,30}\b", validate=_iban_ok),
            # Даты рождения (простая форма)
            PatternSpec("dates", r"\b(0[1-9]|[12]\\d|3[01])\.(0[1-9]|1[0-2])\.(19\\d{2}|20\д{2})\b"),
        ]

        label_overrides = {
            "names": "ФИО",
            "phones": "Телефон",
            "emails": "Email",
            "addresses": "Адрес",
            "documents": "Документ",
            "bank_details": "Банк. реквизит",
            "iban": "IBAN",
            "dates": "Дата",
        }
        for spec in self._specs:
            if spec.label is None and spec.kind in label_overrides:
                spec.label = label_overrides[spec.kind]

        self._specs.extend([
            PatternSpec("badge_numbers", r"\b(?:таб\.?|табельный)\s*(?:номер|№)\s*\\d{3,10}\b", label="Табельный номер"),
            PatternSpec("registration_numbers", r"\b(?:огрн(?:ип)?|грн|рег\.?\s*№?|окпо|оквэд|свид\.?|гос\.?рег\.?№?)\s*[:№-]*[a-z0-9\-]{5,25}\b", label="Регистрационный номер"),
            PatternSpec("domains", r"\b(?:[a-z0-9](?:[a-z0-9-]{0,61}[a-z0-9])?\.)+[a-z]{2,24}\b", label="Домен"),
            PatternSpec("urls", r"\bhttps?://[^\s<'\"]+", label="Ссылка"),
        ])

        self._base_specs: List[PatternSpec] = list(self._specs)

    @staticmethod
    def _normalize_pattern(pattern: str) -> str:
        match = re.match(r'^\(\?([aiLmsux]+)\)(.*)$', pattern, flags=re.DOTALL)
        if match:
            flags, rest = match.groups()
            return f'(?{flags}:{rest})'
        return pattern

    # --------------------------------- API ---------------------------------

    async def process(
        self,
        file_path: str | Path,
        anonymization_mode: str = "replace",  # replace | mask | remove | pseudonym
        exclude_types: List[str] | None = None,
        custom_patterns: List[dict[str, str]] | List[str] | None = None,
        progress_callback: Callable[[dict[str, Any]], Awaitable[None]] | None = None,
        **kwargs,
    ) -> DocumentResult:
        """
        Обезличивание документа.
        """
        async def _notify(stage: str, percent: float, **payload: Any) -> None:
            if not progress_callback:
                return
            data: dict[str, Any] = {"stage": stage, "percent": float(percent)}
            for key, value in payload.items():
                if value is None:
                    continue
                data[key] = value
            try:
                await progress_callback(data)
            except Exception:
                logger.debug("Anonymizer progress callback failed at %s", stage, exc_info=True)

        # 1) Достаём текст
        success, text = await FileFormatHandler.extract_text_from_file(file_path)
        if not success:
            raise ProcessingError(f"Не удалось извлечь текст: {text}", "EXTRACTION_ERROR")

        cleaned_text = TextProcessor.clean_text(text)
        if not cleaned_text.strip():
            raise ProcessingError("Документ не содержит текста", "EMPTY_DOCUMENT")

        words_count = len(cleaned_text.split())
        await _notify("text_extracted", 15, words=words_count)

        # 2) Готовим исключения/карту
        self.anonymization_map = {}
        exclude_set = {item.lower() for item in (exclude_types or [])}

        custom_specs = self._prepare_custom_specs(custom_patterns)

        if path.suffix.lower() == ".docx":
            result_data = await self._anonymize_docx_document(
                path,
                anonymization_mode=anonymization_mode,
                exclude_set=exclude_set,
                custom_specs=custom_specs,
                progress_callback=_notify,
            )
            await _notify("completed", 100, masked=len(self.anonymization_map))
            return DocumentResult.success_result(
                data=result_data, message="Обезличивание документа успешно завершено"
            )

        use_ai = (
            self._ai_enabled
            and self._openai_service is not None
            and not exclude_set
            and not custom_specs
            and anonymization_mode in {"replace", "mask", "remove", "pseudonym"}
        )

        if use_ai:
            await _notify("ai_prepare", 35, mode=anonymization_mode)
            try:
                anonymized_text, report = await self._anonymize_with_ai(
                    cleaned_text,
                    mode=anonymization_mode,
                    progress_hook=_notify,
                )
            except Exception as exc:  # noqa: BLE001
                logger.warning("AI anonymization failed, using rule-based fallback: %s", exc)
                report = {"engine": "fallback", "notes": ["AI anonymization failed; rule-based fallback applied"]}
                use_ai = False
        if not use_ai:
            await _notify("pattern_prepare", 35, excluded=len(exclude_set))
            anonymized_text, rule_report = self._anonymize_text(
                cleaned_text,
                anonymization_mode,
                exclude_set,
                custom_specs=custom_specs,
            )
            report = rule_report if use_ai is False else report
            await _notify("anonymizing", 65, masked=len(self.anonymization_map))

        report["excluded_types"] = sorted(exclude_set) if exclude_set else []
        if custom_specs and "custom_patterns" not in report:
            report["custom_patterns"] = [
                {"kind": spec.kind, "label": spec.label} for spec in custom_specs
            ]

        # 4) Ответ
        await _notify("finalizing", 85, masked=len(self.anonymization_map))

        result_data = {
            "anonymized_text": anonymized_text,
            "anonymization_report": report,
            "anonymization_map": self.anonymization_map.copy(),
            "original_file": str(file_path),
            "mode": anonymization_mode,
        }
        if custom_specs:
            result_data["applied_custom_patterns"] = [
                {"kind": spec.kind, "label": spec.label} for spec in custom_specs
            ]
        await _notify("completed", 100, masked=len(self.anonymization_map))

        return DocumentResult.success_result(
            data=result_data, message="Обезличивание документа успешно завершено"
        )

    # ----------------------------- ЯДРО ЛОГИКИ -----------------------------

    def _anonymize_text(
        self,
        text: str,
        mode: str,
        exclude: set[str] | None = None,
        custom_specs: List[PatternSpec] | None = None,
    ) -> tuple[str, dict[str, Any]]:
        """
        Один проход слева-направо с единым комбинированным паттерном.
        Пользовательские шаблоны подключаются динамически.
        """
        exclude = set(exclude or set())
        specs: List[PatternSpec] = list(self._base_specs)
        if custom_specs:
            specs.extend(custom_specs)

        exclude_custom_all = any(key in exclude for key in ("custom", "custom_patterns"))

        active_specs: List[PatternSpec] = []
        label_map: Dict[str, str] = {}
        for spec in specs:
            if spec.kind in exclude:
                continue
            if exclude_custom_all and spec.kind.startswith("custom_"):
                continue
            active_specs.append(spec)
            label_map[spec.kind] = spec.label or self._pretty_type(spec.kind)

        if not active_specs:
            return text, {"processed_items": [], "statistics": {}, "type_labels": {}}

        stats: Dict[str, int] = {kind: 0 for kind in label_map}
        processed_items: List[Dict[str, Any]] = []

        named_parts = [f"(?P<{s.kind}>{s.pattern})" for s in active_specs]
        combined = re.compile("|".join(named_parts), flags=re.IGNORECASE | re.UNICODE)

        out: List[str] = []
        idx = 0
        counts_by_type: Dict[str, int] = {}

        for match in combined.finditer(text):
            start, end = match.start(), match.end()
            if start < idx:
                continue

            spec = None
            for candidate in active_specs:
                if match.group(candidate.kind):
                    spec = candidate
                    break
            if spec is None:
                continue

            original = match.group(0)

            if spec.validate and not spec.validate(original):
                continue

            out.append(text[idx:start])

            data_type = label_map.get(spec.kind, self._pretty_type(spec.kind))
            replacement = self._get_replacement_deterministic(
                original=original,
                kind=spec.kind,
                data_type=data_type,
                mode=mode,
                counts_by_type=counts_by_type,
            )
            out.append(replacement)
            idx = end

            left = max(0, start - 20)
            right = min(len(text), end + 20)
            processed_items.append(
                {
                    "type": spec.kind,
                    "label": data_type,
                    "value": original,
                    "start": start,
                    "end": end,
                    "snippet": text[left:right],
                }
            )
            stats[spec.kind] = stats.get(spec.kind, 0) + 1

        out.append(text[idx:])
        report = {
            "processed_items": processed_items,
            "statistics": stats,
            "type_labels": label_map,
            "engine": "pattern",
        }
        if custom_specs:
            report["custom_patterns"] = [
                {"kind": spec.kind, "label": label_map.get(spec.kind, spec.label)}
                for spec in custom_specs
            ]
        return "".join(out), report

    def _build_active_specs(
        self,
        exclude: set[str],
        custom_specs: List[PatternSpec] | None = None,
    ) -> tuple[list[PatternSpec], Dict[str, str], re.Pattern[str]]:
        specs: List[PatternSpec] = list(self._base_specs)
        if custom_specs:
            specs.extend(custom_specs)

        exclude_custom_all = any(key in exclude for key in ("custom", "custom_patterns"))

        active_specs: List[PatternSpec] = []
        label_map: Dict[str, str] = {}
        for spec in specs:
            if spec.kind in exclude:
                continue
            if exclude_custom_all and spec.kind.startswith("custom_"):
                continue
            active_specs.append(spec)
            label_map[spec.kind] = spec.label or self._pretty_type(spec.kind)

        if not active_specs:
            combined = re.compile(r"^$", flags=re.IGNORECASE | re.UNICODE)
        else:
            named_parts = [f"(?P<{s.kind}>{s.pattern})" for s in active_specs]
            combined = re.compile("|".join(named_parts), flags=re.IGNORECASE | re.UNICODE)
        return active_specs, label_map, combined

    def _anonymize_block(
        self,
        text: str,
        active_specs: List[PatternSpec],
        label_map: Dict[str, str],
        combined: re.Pattern[str],
        *,
        mode: str,
        *,
        stats: Dict[str, int],
        processed_items: List[Dict[str, Any]],
        counts_by_type: Dict[str, int],
    ) -> tuple[str, List[Dict[str, Any]]]:
        if not active_specs or not text:
            return text, []

        out: List[str] = []
        idx = 0
        matches: List[Dict[str, Any]] = []

        for match in combined.finditer(text):
            start, end = match.start(), match.end()
            if start < idx:
                continue

            spec = None
            for candidate in active_specs:
                if match.group(candidate.kind):
                    spec = candidate
                    break
            if spec is None:
                continue

            original = match.group(0)
            if spec.validate and not spec.validate(original):
                continue

            out.append(text[idx:start])
            data_type = label_map.get(spec.kind, self._pretty_type(spec.kind))
            replacement = self._get_replacement_deterministic(
                original=original,
                kind=spec.kind,
                data_type=data_type,
                mode=mode,
                counts_by_type=counts_by_type,
            )
            out.append(replacement)
            idx = end

            left = max(0, start - 20)
            right = min(len(text), end + 20)
            processed_items.append(
                {
                    "type": spec.kind,
                    "label": data_type,
                    "value": original,
                    "start": start,
                    "end": end,
                    "snippet": text[left:right],
                }
            )
            stats[spec.kind] = stats.get(spec.kind, 0) + 1
            matches.append({"start": start, "end": end, "replacement": replacement})

        out.append(text[idx:])
        return "".join(out), matches

    def _iter_table_paragraphs(self, table: Table):
        for row in table.rows:
            for cell in row.cells:
                yield from self._iter_cell_paragraphs(cell)

    def _iter_cell_paragraphs(self, cell: _Cell):
        for paragraph in cell.paragraphs:
            yield paragraph
        for table in cell.tables:
            yield from self._iter_table_paragraphs(table)

    def _iter_document_paragraphs(self, doc: Document):
        for paragraph in doc.paragraphs:
            yield paragraph
        for table in doc.tables:
            yield from self._iter_table_paragraphs(table)
        for section in doc.sections:
            for part_name in (
                "header",
                "first_page_header",
                "even_page_header",
                "footer",
                "first_page_footer",
                "even_page_footer",
            ):
                part = getattr(section, part_name, None)
                if part is None:
                    continue
                for paragraph in part.paragraphs:
                    yield paragraph
                for table in part.tables:
                    yield from self._iter_table_paragraphs(table)

    def _apply_matches_to_runs(self, runs: List, matches: List[Dict[str, Any]]) -> None:
        if not matches or not runs:
            return

        runs_info: List[Dict[str, Any]] = []
        cursor = 0
        for run in runs:
            text = run.text or ""
            runs_info.append({"run": run, "start": cursor, "end": cursor + len(text)})
            cursor += len(text)

        for match in reversed(matches):
            start = match["start"]
            end = match["end"]
            replacement = match["replacement"]
            affected_indices = [
                idx
                for idx, info in enumerate(runs_info)
                if info["end"] > start and info["start"] < end
            ]
            if not affected_indices:
                continue
            first_idx = affected_indices[0]
            last_idx = affected_indices[-1]

            for idx in affected_indices:
                info = runs_info[idx]
                run = info["run"]
                run_text = run.text or ""
                run_start = info["start"]
                run_end = info["end"]
                local_start = max(0, start - run_start)
                local_end = min(len(run_text), end - run_start)

                if idx == first_idx and idx == last_idx:
                    before = run_text[:local_start]
                    after = run_text[local_end:]
                    run.text = before + replacement + after
                elif idx == first_idx:
                    before = run_text[:local_start]
                    run.text = before + replacement
                elif idx == last_idx:
                    after = run_text[local_end:]
                    run.text = after
                else:
                    run.text = ""

                new_len = len(run.text or "")
                delta = new_len - (run_end - run_start)
                info["end"] = info["start"] + new_len
                for j in range(idx + 1, len(runs_info)):
                    runs_info[j]["start"] += delta
                    runs_info[j]["end"] += delta

    async def _anonymize_docx_document(
        self,
        file_path: Path,
        *,
        anonymization_mode: str,
        exclude_set: set[str],
        custom_specs: List[PatternSpec] | None,
        progress_callback: Callable[[dict[str, Any]], Awaitable[None]] | None,
    ) -> dict[str, Any]:
        doc = Document(str(file_path))

        active_specs, label_map, combined = self._build_active_specs(exclude_set, custom_specs)
        stats: Dict[str, int] = {spec.kind: 0 for spec in active_specs}
        processed_items: List[Dict[str, Any]] = []
        plain_parts: List[str] = []
        counts_by_type: Dict[str, int] = {}

        paragraphs = list(self._iter_document_paragraphs(doc))
        total = len(paragraphs)

        for index, paragraph in enumerate(paragraphs, start=1):
            runs = paragraph.runs
            if not runs:
                plain_parts.append("")
                continue
            original_text = "".join(run.text or "" for run in runs)
            if not original_text:
                plain_parts.append("")
                continue

            anonymized_text, matches = self._anonymize_block(
                original_text,
                active_specs,
                label_map,
                combined,
                mode=anonymization_mode,
                stats=stats,
                processed_items=processed_items,
                counts_by_type=counts_by_type,
            )
            plain_parts.append(anonymized_text)
            if matches:
                self._apply_matches_to_runs(runs, matches)
            if progress_callback:
                percent = 40 + (index / max(1, total)) * 40
                await progress_callback("anonymizing", percent, masked=len(self.anonymization_map))

        output_path = self._build_human_friendly_temp_path(file_path.stem, "анонимизация", ".docx")
        doc.save(str(output_path))

        report = {
            "processed_items": processed_items,
            "statistics": stats,
            "type_labels": label_map,
            "engine": "pattern_docx",
            "excluded_types": sorted(exclude_set) if exclude_set else [],
        }
        if custom_specs:
            report["custom_patterns"] = [{"kind": spec.kind, "label": label_map.get(spec.kind, spec.label)} for spec in custom_specs]

        anonymized_text = "\n\n".join(part for part in plain_parts if part is not None)

        return {
            "anonymized_text": anonymized_text.strip(),
            "anonymization_report": report,
            "anonymization_map": self.anonymization_map.copy(),
            "original_file": str(file_path),
            "mode": anonymization_mode,
            "anonymized_docx": str(output_path),
        }

    def _prepare_custom_specs(
        self, custom_patterns: List[dict[str, str]] | List[str] | None
    ) -> List[PatternSpec]:
        specs: List[PatternSpec] = []
        if not custom_patterns:
            return specs

        used_kinds = {spec.kind for spec in self._base_specs}

        for idx, entry in enumerate(custom_patterns, start=1):
            pattern = ""
            label = ""
            if isinstance(entry, dict):
                pattern = str(entry.get("pattern", "")).strip()
                label = str(entry.get("name") or entry.get("label") or "").strip()
            else:
                pattern = str(entry).strip()

            if not pattern:
                continue

            normalized_pattern = self._normalize_pattern(pattern)
            try:
                re.compile(normalized_pattern, re.IGNORECASE | re.UNICODE)
            except re.error as exc:
                logger.warning("Skip custom pattern %s: %s", pattern, exc)
                continue

            slug_source = label or f"pattern_{idx}"
            slug = re.sub(r"[^a-z0-9]+", "_", slug_source.lower()).strip("_")
            if not slug:
                slug = f"pattern_{idx}"

            base_kind = f"custom_{slug}"
            kind = base_kind
            counter = 2
            while kind in used_kinds:
                kind = f"{base_kind}_{counter}"
                counter += 1
            used_kinds.add(kind)

            human_label = label or "Пользовательский шаблон"
            specs.append(PatternSpec(kind, normalized_pattern, label=human_label))

        return specs

    async def _anonymize_with_ai(
        self,
        text: str,
        *,
        mode: str,
        progress_hook: Callable[..., Awaitable[None]] | None = None,
    ) -> tuple[str, dict[str, Any]]:
        if self._openai_service is None:
            raise ProcessingError("AI anonymization unavailable", "AI_DISABLED")

        chunks = TextProcessor.split_into_chunks(
            text,
            max_chunk_size=self._ai_chunk_chars,
            overlap=self._ai_chunk_overlap,
        )
        if not chunks:
            return "", {
                "processed_items": [],
                "statistics": {"chunks": 0},
                "type_labels": {},
                "engine": "openai",
                "notes": ["Документ не содержит текста для анонимизации."],
            }

        anonymized_parts: List[str] = []
        processed_items: List[Dict[str, Any]] = []
        stats: Dict[str, int] = {}
        type_labels: Dict[str, str] = {}
        notes_accum: List[str] = []
        replacements_seen: set[tuple[str, str]] = set()
        structured_chunks = 0
        total = len(chunks)

        for index, chunk in enumerate(chunks, start=1):
            if progress_hook:
                percent = 35 + (index / total) * 40
                await progress_hook("ai_chunk", percent, chunk=index, total=total)

            user_prompt = self._build_ai_prompt(
                chunk=chunk,
                mode=mode,
                index=index,
                total=total,
            )

            response = await self._openai_service.ask_legal(
                system_prompt=_AI_SYSTEM_PROMPT,
                user_text=user_prompt,
                force_refresh=True,
            )

            if not response.get("ok"):
                raise ProcessingError(
                    response.get("error") or "AI anonymization failed",
                    "AI_ERROR",
                )

            chunk_text = response.get("text") or ""
            parsed = self._parse_ai_json_output(chunk_text)

            replacements: List[Dict[str, Any]] = []
            if isinstance(parsed, dict):
                structured_chunks += 1
                chunk_anonymized = str(
                    parsed.get("anonymized_chunk")
                    or parsed.get("anonymized_text")
                    or parsed.get("text")
                    or ""
                ).strip()
                notes = parsed.get("notes") or []
                for note in notes:
                    note_str = str(note or "").strip()
                    if note_str:
                        notes_accum.append(note_str)
                raw_replacements = parsed.get("replacements") or []
                if isinstance(raw_replacements, list):
                    replacements = [item for item in raw_replacements if isinstance(item, dict)]
                if not chunk_anonymized:
                    chunk_anonymized = self._strip_html(chunk_text).strip()
            else:
                chunk_anonymized = self._strip_html(chunk_text).strip()

            anonymized_parts.append(chunk_anonymized)

            for item in replacements:
                original = str(item.get("original") or "").strip()
                replacement = str(item.get("replacement") or "").strip()
                if not original:
                    continue

                kind_raw = str(item.get("kind") or "").strip().lower()
                kind = kind_raw or "custom"
                label = str(item.get("label") or "").strip() or self._pretty_type(kind)
                snippet = str(item.get("context") or "").strip()

                key = (original.lower(), replacement.lower())
                if key in replacements_seen:
                    continue
                replacements_seen.add(key)

                stats[kind] = stats.get(kind, 0) + 1
                if label:
                    type_labels.setdefault(kind, label)

                processed_item: Dict[str, Any] = {
                    "type": kind,
                    "label": label,
                    "value": original,
                }
                if snippet:
                    processed_item["snippet"] = snippet
                if replacement:
                    processed_item["replacement"] = replacement

                processed_items.append(processed_item)
                if replacement:
                    self.anonymization_map.setdefault(original, replacement)
                else:
                    self.anonymization_map.setdefault(original, "")

        anonymized = "\n\n".join(part for part in anonymized_parts if part)

        report_stats: Dict[str, Any] = {"chunks": total}
        for key, value in stats.items():
            report_stats[key] = value

        if notes_accum:
            unique_notes: List[str] = []
            seen_notes_set: set[str] = set()
            for note in notes_accum:
                if note in seen_notes_set:
                    continue
                seen_notes_set.add(note)
                unique_notes.append(note)
                if len(unique_notes) >= 5:
                    break
            notes_accum = unique_notes

        if structured_chunks:
            report_notes = notes_accum
        else:
            report_notes = notes_accum or [
                "Структуру замен получить не удалось.",
            ]

        report = {
            "processed_items": processed_items,
            "statistics": report_stats,
            "type_labels": type_labels,
            "engine": "openai",
            "notes": report_notes,
        }
        return anonymized, report

    @staticmethod
    def _build_ai_prompt(*, chunk: str, mode: str, index: int, total: int) -> str:
        mode = mode.lower().strip()
        instructions = {
            "replace": (
                "Замени все персональные данные универсальными маркерами вида [PERSON], [ORG], [PHONE], [ADDRESS],"
                " сохраняя смысл текста."
            ),
            "remove": (
                "Полностью удаляй персональные данные, оставляя маркер [REDACTED] на их месте, чтобы указать на удаление."
            ),
            "mask": (
                "Замаскируй персональные данные символами *, сохранив формат (например, номера телефонов и документов)."
            ),
            "pseudonym": (
                "Заменяй каждую сущность на псевдоним вида PERSON_1, COMPANY_1, LOCATION_1. Используй один и тот же псевдоним для повторений внутри фрагмента."
            ),
        }
        instruction = instructions.get(mode, instructions["replace"])

        schema = dedent(
            """
            Верни строго JSON без пояснений и форматирующего текста. Структура:
            {
              "anonymized_chunk": "...",
              "replacements": [
                {
                  "kind": "names|phones|emails|addresses|documents|bank_details|iban|dates|badge_numbers|registration_numbers|domains|urls|custom",
                  "label": "Человекочитаемое название типа",
                  "original": "Текст из исходного фрагмента до замены (с сохранением регистра и знаков)",
                  "replacement": "Текст, который ты поставил вместо оригинала (пустая строка, если удалил)",
                  "context": "Короткий (<=120 символов) фрагмент исходного текста вокруг замены"
                }
              ],
              "notes": ["опционально пояснения"]
            }
            Если замен нет, верни пустой массив replacements. Используй ровно те же псевдонимы, маски
            или маркеры, которые вставляешь в текст. Не добавляй Markdown и комментарии.
            """
        ).strip()

        return dedent(
            f"""
            Ты выступаешь сервисом анонимизации. {schema}
            Перед тобой часть документа ({index} из {total}). {instruction}
            Сохраняй структуру, нумерацию, правовые ссылки и абзацы.
            Работай аккуратно: не удаляй лишних символов и не меняй порядок абзацев.

            <НАЧАЛО_ФРАГМЕНТА>
            {chunk}
            <КОНЕЦ_ФРАГМЕНТА>
            """
        ).strip()

    @staticmethod
    def _strip_html(content: str) -> str:
        if not content:
            return ""
        return re.sub(r"<[^>]+>", "", content)

    @staticmethod
    def _parse_ai_json_output(content: str) -> Dict[str, Any] | None:
        if not content:
            return None

        cleaned = content.strip()
        if not cleaned:
            return None

        fence_match = re.search(r"```(?:json)?\s*(\{.*\})\s*```", cleaned, re.DOTALL)
        if fence_match:
            cleaned = fence_match.group(1).strip()

        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            start = cleaned.find("{")
            end = cleaned.rfind("}")
            if start != -1 and end != -1 and end > start:
                snippet = cleaned[start : end + 1]
                try:
                    return json.loads(snippet)
                except json.JSONDecodeError as exc:  # noqa: PERF203
                    logger.debug("Failed to parse AI anonymization JSON snippet: %s", snippet, exc_info=True)
                    return None
            logger.debug("Failed to parse AI anonymization JSON: %s", cleaned, exc_info=True)
            return None

    # ------------------------- ПОДМЕНА / ФОРМАТИРОВАНИЕ -------------------------

    @staticmethod
    def _pretty_type(kind: str) -> str:
        mapping = {
            "names": "Лицо",
            "phones": "Телефон",
            "emails": "Email",
            "addresses": "Адрес",
            "documents": "Документ",
            "bank_details": "БанкРеквизит",
            "iban": "IBAN",
            "dates": "Дата",
            "badge_numbers": "Табельный номер",
            "registration_numbers": "Регистрационный номер",
            "domains": "Домен",
            "urls": "Ссылка",
        }
        if kind.startswith("custom_"):
            return "Пользовательский шаблон"
        return mapping.get(kind, "Данные")

    def _get_replacement_deterministic(
        self,
        *,
        original: str,
        kind: str,
        data_type: str,
        mode: str,
        counts_by_type: Dict[str, int],
    ) -> str:
        """
        remove   → [УДАЛЕНО]
        mask     → сохраняем формат (телефон/email — спец. маска)
        replace  → [Тип-N] (инкремент по типу)
        pseudonym→ [Тип~HMAC] (детерминированный по секрету ANON_SECRET)
        """
        if original in self.anonymization_map:
            return self.anonymization_map[original]

        if mode == "remove":
            replacement = "[УДАЛЕНО]"
        elif mode == "mask":
            if kind == "phones":
                replacement = _mask_phone(original)
            elif kind == "emails":
                replacement = _mask_email(original)
            else:
                replacement = _mask_preserve(original)
        elif mode == "pseudonym":
            secret_bytes = self._secret.encode("utf-8") if self._secret else b"default-anon-secret"
            secret = secret_bytes
            code = _pseudo_id(kind, original, secret)
            replacement = f"[{data_type}~{code}]"
        else:  # replace
            n = counts_by_type.get(data_type, 0) + 1
            counts_by_type[data_type] = n
            replacement = f"[{data_type}-{n}]"

        self.anonymization_map[original] = replacement
        return replacement

    # -------------------------------- ВОССТАНОВЛЕНИЕ ----------------------------
