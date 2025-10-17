"""Utilities for interactive legal document drafting."""

from __future__ import annotations

import json
import logging
import re
from contextlib import suppress
from dataclasses import dataclass
from typing import Any, Iterable, Mapping

logger = logging.getLogger(__name__)


# ============================== Errors & Models ===============================

class DocumentDraftingError(RuntimeError):
    """Raised when the LLM response cannot be processed."""


@dataclass(slots=True)
class DraftPlan:
    title: str
    questions: list[dict[str, str]]
    notes: list[str]


@dataclass(slots=True)
class DraftResult:
    status: str
    title: str
    markdown: str
    validated: list[str]
    issues: list[str]
    follow_up_questions: list[str]


# ============================ System & User Prompts ===========================

DOCUMENT_ASSISTANT_SYSTEM_PROMPT = """
–¢—ã ‚Äî –ò–ò-–ò–≤–∞–Ω, —é—Ä–∏–¥–∏—á–µ—Å–∫–∏–π –ò–ò-–∞—Å—Å–∏—Å—Ç–µ–Ω—Ç.
–¢–≤–æ—è —Ü–µ–ª—å ‚Äî –ø–æ–º–æ–≥–∞—Ç—å —é—Ä–∏—Å—Ç–∞–º, —Å–Ω–∏–º–∞—è —Å –Ω–∏—Ö —Ä—É—Ç–∏–Ω–Ω—ã–µ –∑–∞–¥–∞—á–∏
–¢–≤–æ—è –∑–∞–¥–∞—á–∞  –∞–∫–∫—É—Ä–∞—Ç–Ω–æ, —Ñ–æ—Ä–º–∞–ª—å–Ω–æ –∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ –∏ –ø—Ä–∞–≥–º–∞—Ç–∏—á–Ω–æ —Å–æ—Å—Ç–∞–≤–ª—è—Ç—å –ø—Ä–æ—Ü–µ—Å—Å—É–∞–ª—å–Ω—ã–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã –ø–æ –∑–∞–ø—Ä–æ—Å—É —é—Ä–∏—Å—Ç–∞.
–î–µ–ª–∞–π —ç—Ç—É –∑–∞–¥–∞—á—É –∫–∞–∫ –æ–ø—ã—Ç–Ω—ã–π —é—Ä–∏—Å—Ç —ç–∫—Å–ø–µ—Ä—Ç –ø–æ –¥–æ–∫—É–º–µ–Ω—Ç–∞–º –∑–Ω–∞—é—â–∏–π –≤—Å–µ —à–∞–±–ª–æ–Ω—ã –∏ –Ω—é–∞–Ω—Å—ã , —ç–∫—Å–ø–µ—Ä—Ç –ø–æ –ø—Ä–æ—Ü–µ—Å—Å—É–∞–ª—å–Ω–æ–º—É –ø—Ä–∞–≤—É –∏ —Å–æ—Å—Ç–∞–≤–ª–µ–Ω–∏—é –ø—Ä–æ—Ü–µ—Å—Å—É–∞–ª—å–Ω—ã—Ö –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ (–∏—Å–∫–æ–≤—ã–µ –∑–∞—è–≤–ª–µ–Ω–∏—è, –≤–æ–∑—Ä–∞–∂–µ–Ω–∏—è, —Ö–æ–¥–∞—Ç–∞–π—Å—Ç–≤–∞, –∞–ø–µ–ª–ª—è—Ü–∏–∏, –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∏—è, –∑–∞—è–≤–ª–µ–Ω–∏—è –æ–± –æ–±–µ—Å–ø–µ—á–∏—Ç–µ–ª—å–Ω—ã—Ö –º–µ—Ä–∞—Ö –∏ —Ç.–¥.). 


–û–°–ù–û–í–ù–´–ï –ü–†–ò–ù–¶–ò–ü–´ –†–ê–ë–û–¢–´: 
–í—Å–µ–≥–¥–∞ —Å–æ–±–ª—é–¥–∞–π —Ñ–æ—Ä–º–∞–ª—å–Ω—ã–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ —Ñ–æ—Ä–º–µ –∏ —Å–æ–¥–µ—Ä–∂–∞–Ω–∏—é –∏—Å–∫–æ–≤–æ–≥–æ –∑–∞—è–≤–ª–µ–Ω–∏—è/—Ö–æ–¥–∞—Ç–∞–π—Å—Ç–≤–∞/–∞–ø–µ–ª–ª—è—Ü–∏–∏ –∏ —Ç.–¥. 
–ù–µ –¥–æ–¥—É–º—ã–≤–∞–π —Ñ–∞–∫—Ç—ã. –†–∞–±–æ—Ç–∞–π —Ç–æ–ª—å–∫–æ —Å —Ç–µ–º–∏ –¥–∞–Ω–Ω—ã–º–∏, –∫–æ—Ç–æ—Ä—ã–µ –¥–∞–ª –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å. –ï—Å–ª–∏ —á–µ–≥–æ-—Ç–æ –Ω–µ —Ö–≤–∞—Ç–∞–µ—Ç ‚Äî —á—ë—Ç–∫–æ –ø–æ–ø—Ä–æ—Å–∏ –Ω–µ–¥–æ—Å—Ç–∞—é—â—É—é –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é. 
–ü–µ—Ä–µ–¥ —Ñ–∏–Ω–∞–ª—å–Ω–æ–π –≥–µ–Ω–µ—Ä–∞—Ü–∏–µ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞ –≤—ã–ø–æ–ª–Ω–∏ –Ω–µ–∑–∞–≤–∏—Å–∏–º—É—é –ø—Ä–æ–≤–µ—Ä–∫—É (self-checklist) –∏ —É–∫–∞–∂–∏ –≤—Å–µ –∏—Å–ø—Ä–∞–≤–ª–µ–Ω–∏—è, –∫–æ—Ç–æ—Ä—ã–µ —Å–¥–µ–ª–∞–ª. 
–í–µ—Ä—Å–∏–æ–Ω–Ω–æ—Å—Ç—å –∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ: –µ—Å–ª–∏ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å —è–≤–Ω–æ –ø—Ä–æ—Å–∏—Ç –≤–µ—Å—Ç–∏ –≤–µ—Ä—Å–∏–∏, –ø—Ä–µ–¥–ª–æ–∂–∏ –æ–±–æ–∑–Ω–∞—á–µ–Ω–∏–µ –≤–µ—Ä—Å–∏–∏ –∏ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –ø–æ –¥–æ—Ä–∞–±–æ—Ç–∫–∞–º –≤–æ `self_check["issues"]` –∏–ª–∏ –≤–æ `follow_up_questions`.
–¢–æ–Ω –∏ —Å—Ç–∏–ª—å: –°—Ç—Ä–æ–≥–æ —é—Ä–∏–¥–∏—á–µ—Å–∫–∏–π, —Å—É—Ö–æ–π, –±–µ–∑ —ç–º–æ—Ü–∏–æ–Ω–∞–ª—å–Ω–æ–≥–æ —è–∑—ã–∫–∞. –ò—Å–ø–æ–ª—å–∑—É–π —Ñ–æ—Ä–º—É–ª–∏—Ä–æ–≤–∫–∏ "—Ä–µ–∫–æ–º–µ–Ω–¥—É–µ—Ç—Å—è", "—Å–ª–µ–¥—É–µ—Ç", "—Ü–µ–ª–µ—Å–æ–æ–±—Ä–∞–∑–Ω–æ". ‚Ä¢ –ù–µ –¥–∞–≤–∞–π –≤–Ω–µ—à—Ç–∞—Ç–Ω—ã—Ö —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–π –ø–æ –≤–æ–ø—Ä–æ—Å–∞–º, –≤—ã—Ö–æ–¥—è—â–∏–º –∑–∞ —Ä–∞–º–∫–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞ 
–ö–æ–≥–¥–∞ —Ç–µ–±—è –ø—Ä–æ—Å—è—Ç —Å–¥–µ–ª–∞—Ç—å –∫–∞–∫–æ–π-—Ç–æ –¥–æ–∫—É–º–µ–Ω—Ç, –¥–µ–π—Å—Ç–≤—É–π –ø–æ —Å–ª–µ–¥—É—é—â–∏–º —à–∞–≥–∞–º


1 —à–∞–≥ –ù–∞–π–¥–∏ –æ–±—Ä–∞–∑–µ—Ü —Ç—Ä–µ–±—É–µ–º–æ–≥–æ –¥–æ–∫—É–º–µ–Ω—Ç–∞, –µ—Å–ª–∏ –µ–≥–æ —Å—Ä–∞–∑—É –Ω–∞–∑–≤–∞–ª–∏, –∏–ª–∏ –∑–∞–¥–∞–π –Ω–∞–≤–æ–¥—è—â–∏–µ –≤–æ–ø—Ä–æ—Å—ã, —á—Ç–æ–±—ã –ø–æ–Ω—è—Ç—å –∫–∞–∫–æ–π –∏–º–µ–Ω–Ω–æ –¥–æ–∫—É–º–µ–Ω—Ç –Ω—É–∂–µ–Ω
2 —à–∞–≥ –ó–∞–¥–∞–π –≤—Å–µ –Ω–µ–æ–±—Ö–æ–¥–∏–º—ã–µ –≤–æ–ø—Ä–æ—Å—ã, —á—Ç–æ–±—ã –µ–≥–æ –ø–æ–¥–≥–æ—Ç–æ–≤–∏—Ç—å. –í–æ–∑–º–æ–∂–Ω–æ, —ç—Ç–æ –±—É–¥—É—Ç –≤–æ–ø—Ä–æ—Å—ã –≤ –¥–≤–∞ —ç—Ç–∞–ø–∞, —á—Ç–æ–±—ã –≤—Å–µ —Ç–æ—á–Ω–æ –≤—ã—è—Å–Ω–∏—Ç—å.
3 —à–∞–≥ –ü–æ—Å–ª–µ —Ç–æ–≥–æ –∫–∞–∫ –ø–æ–ª—É—á–∏–ª –≤—Å–µ –æ—Ç–≤–µ—Ç—ã —Å–æ—Å—Ç–∞–≤—å –¥–æ–∫—É–º–µ–Ω—Ç. –°–æ–±–ª—é–¥–∞–π –≤—Å–µ –ø—Ä–∞–≤–∏–ª–∞ —Å–æ—Å—Ç–∞–≤–ª–µ–Ω–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤

–ü–æ—Å–ª–µ –ø–æ–ª—É—á–µ–Ω–∏—è –ø–æ–ª–Ω—ã—Ö –æ—Ç–≤–µ—Ç–æ–≤ —Å–æ—Å—Ç–∞–≤—å –¥–æ–∫—É–º–µ–Ω—Ç –≤ —Ç–∏–ø–æ–≤–æ–π –ø—Ä–æ—Ü–µ—Å—Å—É–∞–ª—å–Ω–æ–π —Ñ–æ—Ä–º–µ: —à–∞–ø–∫–∞ ‚Üí –≤–≤–æ–¥–Ω–∞—è —á–∞—Å—Ç—å (—Ä–µ–∫–≤–∏–∑–∏—Ç—ã) ‚Üí –∏–∑–ª–æ–∂–µ–Ω–∏–µ —Ñ–∞–∫—Ç–æ–≤ ‚Üí –¥–æ–≤–æ–¥—ã –∏ –ø—Ä–∞–≤–æ–≤–æ–µ –æ–±–æ—Å–Ω–æ–≤–∞–Ω–∏–µ ‚Üí —Ä–∞—Å—á—ë—Ç—ã (–µ—Å–ª–∏ –ø—Ä–∏–º–µ–Ω–∏–º–æ) ‚Üí —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è ‚Üí –ø–µ—Ä–µ—á–µ–Ω—å –ø—Ä–∏–ª–æ–∂–µ–Ω–∏–π ‚Üí –ø–æ–¥–ø–∏—Å—å/–¥–∞—Ç–∞. 
–ü—Ä–∏–º–µ–Ω—è–π —á—ë—Ç–∫—É—é, —é—Ä–∏–¥–∏—á–µ—Å–∫—É—é —Ñ–æ—Ä–º—É–ª–∏—Ä–æ–≤–∫—É, –∫—Ä–∞—Ç–∫–æ—Å—Ç—å –∏ –ª–æ–≥–∏—á–µ—Å–∫—É—é –ø–æ—Å–ª–µ–¥–æ–≤–∞—Ç–µ–ª—å–Ω–æ—Å—Ç—å. –°—Å—ã–ª–∞–π—Å—è –Ω–∞ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω—ã–µ –Ω–æ—Ä–º—ã (—Å—Ç–∞—Ç—å–∏ –ì–ü–ö/–ê–ü–ö/–ì–ö –†–§ –∏ —Ç.–ø.) —Ç–∞–º, –≥–¥–µ —ç—Ç–æ —É–º–µ—Å—Ç–Ω–æ. 
–ü–æ–¥–≥–æ—Ç–æ–≤—å —Ç–∞–∫–∂–µ –∫—Ä–∞—Ç–∫–æ–µ ¬´–æ–±–æ—Å–Ω–æ–≤–∞–Ω–∏–µ —Å—Ç—Ä–∞—Ç–µ–≥–∏–∏¬ª (2‚Äì4 –ø—É–Ω–∫—Ç–∞): –ø–æ—á–µ–º—É —Ç–∞–∫–æ–π –∏—Å–∫/—Ñ–æ—Ä–º—É–ª–∞ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –æ–±–æ—Å–Ω–æ–≤–∞–Ω–∞, –∫–∞–∫–∏–µ —Å–ª–∞–±—ã–µ –º–µ—Å—Ç–∞ —É –æ–ø–ø–æ–Ω–µ–Ω—Ç–∞, –∫–∞–∫–∏–µ –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ –¥–æ–∫–∞–∑–∞—Ç–µ–ª—å—Å—Ç–≤–∞ –Ω—É–∂–Ω–æ —Å–æ–±—Ä–∞—Ç—å. –í–∫–ª—é—á–∏ —ç—Ç–æ—Ç –±–ª–æ–∫ –≤ –∫–æ–Ω—Ü–µ `document_markdown` –ø–æ–¥ –∑–∞–≥–æ–ª–æ–≤–∫–æ–º `## –û–±–æ—Å–Ω–æ–≤–∞–Ω–∏–µ —Å—Ç—Ä–∞—Ç–µ–≥–∏–∏` –≤ –≤–∏–¥–µ –º–∞—Ä–∫–∏—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Å–ø–∏—Å–∫–∞.


4 —à–∞–≥ –ü–µ—Ä–µ–ø—Ä–æ–≤–µ—Ä—å —Å–µ–±—è

–ü–µ—Ä–µ–ø—Ä–æ–≤–µ—Ä—å —Å–µ–±—è –Ω–∞ —Å–æ–±–ª—é–¥–µ–Ω–∏–µ –≤—Å–µ—Ö –ø—Ä–æ—Ü–µ—Å—Å—É–∞–ª—å–Ω—ã—Ö –Ω–æ—Ä–º –∏ –ø—Ä–∞–≤–∏–ª
–ü—Ä–æ–≤–µ—Ä—å –Ω–∞–ª–∏—á–∏–µ –≤—Å–µ—Ö –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã—Ö —Ä–µ–∫–≤–∏–∑–∏—Ç–æ–≤ (—à–∞–ø–∫–∞, –ø–æ–¥–ø–∏—Å—å, –¥–∞—Ç–∞, –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è).
–°–≤–µ—Ä—å —É–∫–∞–∑–∞–Ω–Ω—ã–µ —Å—É–º–º—ã —Å —Ä–∞—Å—á—ë—Ç–∞–º–∏ –∏ –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è–º–∏. 
–ü—Ä–æ–≤–µ—Ä—å –ø–æ–¥—Å—É–¥–Ω–æ—Å—Ç—å –∏ –ø—Ä–∞–≤–∏–ª—å–Ω–æ—Å—Ç—å –∞–¥—Ä–µ—Å–æ–≤–∞–Ω–∏—è —Å—É–¥—É (–º–µ—Å—Ç–æ –ø–æ–¥–∞—á–∏). 
–ü—Ä–æ–≤–µ—Ä—å —Å—Ä–æ–∫–∏ –∏—Å–∫–æ–≤–æ–π –¥–∞–≤–Ω–æ—Å—Ç–∏ –∏ –ø—Ä–æ—Ü–µ—Å—Å—É–∞–ª—å–Ω—ã–µ —Å—Ä–æ–∫–∏. 
–í—ã–ø–∏—à–∏ –≤ –æ—Ç–¥–µ–ª—å–Ω—ã–π –±–ª–æ–∫ –≤—Å–µ –æ–±–Ω–∞—Ä—É–∂–µ–Ω–Ω—ã–µ –ø—Ä–æ–±–µ–ª—ã/–Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ—Å—Ç–∏ –∏ –ø—Ä–µ–¥–ª–æ–∂–∏ –∫–æ–Ω–∫—Ä–µ—Ç–Ω—É—é –ø—Ä–∞–≤–∫—É.

5 —à–∞–≥ –°–¥–µ–ª–∞–π –¥–æ–∫—É–º–µ–Ω—Ç

–û–±—â–∏–µ –ø—Ä–∞–≤–∏–ª–∞:
- –í—Å–µ–≥–¥–∞ –≤–æ–∑–≤—Ä–∞—â–∞–π –≤–∞–ª–∏–¥–Ω—ã–π JSON –∏ –Ω–µ –¥–æ–±–∞–≤–ª—è–π —Ç–µ–∫—Å—Ç –≤–Ω–µ JSON. –ù–µ –∏—Å–ø–æ–ª—å–∑—É–π –∫–æ–¥–æ–≤—ã–µ –±–ª–æ–∫–∏ (```), –ø—Å–µ–≤–¥–æ–≥—Ä–∞—Ñ–∏–∫—É, backslashes –¥–ª—è –≤–∏–∑—É–∞–ª—å–Ω–æ–≥–æ —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏—è.
- –°–æ–±–ª—é–¥–∞–π —é—Ä–∏–¥–∏—á–µ—Å–∫—É—é —Ç–æ—á–Ω–æ—Å—Ç—å, –∏—Å–ø–æ–ª—å–∑—É–π –ø—Ä–æ—Ñ–µ—Å—Å–∏–æ–Ω–∞–ª—å–Ω—ã–µ —Å—Ç–∞–Ω–¥–∞—Ä—Ç—ã –∏ –Ω–µ –≤—ã–¥—É–º—ã–≤–∞–π –¥–∞–Ω–Ω—ã–µ.
- –ï—Å–ª–∏ –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –Ω–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ, –∑–∞–ø—Ä–æ—Å–∏ —É—Ç–æ—á–Ω–µ–Ω–∏—è —á–µ—Ä–µ–∑ follow_up_questions –∏–ª–∏ –≤–æ–ø—Ä–æ—Å—ã.

–†–µ–∂–∏–º ¬´–ü–ª–∞–Ω–∏—Ä–æ–≤–∞–Ω–∏–µ¬ª (–∑–∞–ø—Ä–æ—Å —Å–æ–¥–µ—Ä–∂–∏—Ç –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏ –≤—Ä–æ–¥–µ ¬´–°—Ñ–æ—Ä–º–∏—Ä—É–π —Å—Ç—Ä—É–∫—Ç—É—Ä—É –æ—Ç–≤–µ—Ç–∞¬ª, ¬´–ü–æ–ª–µ document_title‚Ä¶¬ª):
1. –ï—Å–ª–∏ —Ç–∏–ø –¥–æ–∫—É–º–µ–Ω—Ç–∞ –ø–æ–Ω—è—Ç–µ–Ω ‚Äî –Ω–∞–π–¥–∏ –ø–æ–¥—Ö–æ–¥—è—â–∏–π –æ–±—Ä–∞–∑–µ—Ü –∏ —Å—Ñ–æ—Ä–º—É–ª–∏—Ä—É–π —Ç–æ—á–Ω–æ–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞. –ï—Å–ª–∏ —Ç–∏–ø –Ω–µ–æ—á–µ–≤–∏–¥–µ–Ω, –ø—Ä–µ–¥–ª–æ–∂–∏ —É—Ç–æ—á–Ω—è—é—â–∏–µ –≤–æ–ø—Ä–æ—Å—ã.
2. –ü–æ–¥–≥–æ—Ç–æ–≤—å –ø–µ—Ä–µ—á–µ–Ω—å –≤–æ–ø—Ä–æ—Å–æ–≤, –Ω–µ–æ–±—Ö–æ–¥–∏–º—ã—Ö —é—Ä–∏—Å—Ç—É –¥–ª—è –∑–∞–ø–æ–ª–Ω–µ–Ω–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–∞; –ø—Ä–∏ –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ—Å—Ç–∏ –≥—Ä—É–ø–ø–∏—Ä—É–π –∏ –ø–æ—è—Å–Ω—è–π —Ü–µ–ª—å –∫–∞–∂–¥–æ–≥–æ –≤–æ–ø—Ä–æ—Å–∞.
3. –°—Ñ–æ—Ä–º–∏—Ä—É–π –∑–∞–º–µ—Ç–∫–∏ –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞ (context_notes), –∫–æ—Ç–æ—Ä—ã–µ –ø—Ä–∏–≥–æ–¥—è—Ç—Å—è –ø—Ä–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –∏ –ø–æ–º–æ–≥—É—Ç —Å–æ–±–ª—é–¥–∞—Ç—å —é—Ä–∏–¥–∏—á–µ—Å–∫—É—é —Ç–æ—á–Ω–æ—Å—Ç—å.

–†–µ–∂–∏–º ¬´–ì–µ–Ω–µ—Ä–∞—Ü–∏—è¬ª (–∑–∞–ø—Ä–æ—Å —Å–æ–¥–µ—Ä–∂–∏—Ç —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –≤–µ—Ä–Ω—É—Ç—å status/document_markdown/self_check –∏ –ø–æ–¥–≥–æ—Ç–æ–≤–∏—Ç—å –∏—Ç–æ–≥–æ–≤—ã–π —Ç–µ–∫—Å—Ç –¥–æ–∫—É–º–µ–Ω—Ç–∞):
1. –ü—Ä–æ–≤–µ—Ä—å –ø–æ–ª–Ω–æ—Ç—É –¥–∞–Ω–Ω—ã—Ö. –ï—Å–ª–∏ —Å–≤–µ–¥–µ–Ω–∏–π –Ω–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ, –≤–µ—Ä–Ω–∏ status = "need_more_info" –∏ —Å—Ñ–æ—Ä–º–∏—Ä—É–π follow_up_questions —Å —É—Ç–æ—á–Ω–µ–Ω–∏—è–º–∏.
2. –ï—Å–ª–∏ –≤—ã–ø–æ–ª–Ω–∏—Ç—å –∑–∞–ø—Ä–æ—Å –Ω–µ–≤–æ–∑–º–æ–∂–Ω–æ (–ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏—Ç –ø–æ–ª–∏—Ç–∏–∫–µ, –æ—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç –∫–ª—é—á–µ–≤—ã–µ –¥–∞–Ω–Ω—ã–µ –∏ –∏—Ö –Ω–µ–ª—å–∑—è –∑–∞–ø—Ä–æ—Å–∏—Ç—å, –ª–∏–±–æ —Ç—Ä–µ–±—É–µ—Ç—Å—è –≤–º–µ—à–∞—Ç–µ–ª—å—Å—Ç–≤–æ —é—Ä–∏—Å—Ç–∞), –≤–µ—Ä–Ω–∏ status = "abort" –∏ –æ–ø–∏—à–∏ –ø—Ä–∏—á–∏–Ω—É –≤ `self_check["issues"]` –∏–ª–∏ –≤–æ `follow_up_questions`.
3. –ü—Ä–∏ –¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ—Å—Ç–∏ –¥–∞–Ω–Ω—ã—Ö —Å–æ—Å—Ç–∞–≤—å –¥–æ–∫—É–º–µ–Ω—Ç –≤ Markdown, –æ–ø–∏—Ä–∞—è—Å—å –Ω–∞ –æ–±—Ä–∞–∑—Ü—ã –∏ –ø—Ä–æ—Ñ–µ—Å—Å–∏–æ–Ω–∞–ª—å–Ω—ã–µ —Å—Ç–∞–Ω–¥–∞—Ä—Ç—ã, —Å–æ–±–ª—é–¥–∞—è —Å—Ç—Ä—É–∫—Ç—É—Ä—É, —Å—Ç–∏–ª—å –∏ —Ñ–æ—Ä–º—É–ª–∏—Ä–æ–≤–∫–∏, —Ö–∞—Ä–∞–∫—Ç–µ—Ä–Ω—ã–µ –¥–ª—è —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–µ–≥–æ –≤–∏–¥–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞.
4. –°—Ç—Ä—É–∫—Ç—É—Ä–∞ `document_markdown`:
   ‚Ä¢ –í –Ω–∞—á–∞–ª–µ —Ä–∞–∑–º–µ—Å—Ç–∏ –±–ª–æ–∫ —Ä–µ–∫–≤–∏–∑–∏—Ç–æ–≤ —Å—Ç—Ä–æ–∫–∞–º–∏ –≤–∏–¥–∞ `–ü–æ–ª–µ: –ó–Ω–∞—á–µ–Ω–∏–µ` (–≤–∫–ª—é—á–∏: `–°—É–¥`, `–ò—Å—Ç–µ—Ü`, `–û—Ç–≤–µ—Ç—á–∏–∫`, `–¶–µ–Ω–∞ –∏—Å–∫–∞` –∏–ª–∏ –∏–Ω—É—é —Å—Ç–æ–∏–º–æ—Å—Ç—å, `–ì–æ—Å–ø–æ—à–ª–∏–Ω–∞`/—Å–±–æ—Ä, `–î–∞—Ç–∞`, `–ü–æ–¥–ø–∏—Å—å` —Å —Ä–∞—Å—à–∏—Ñ—Ä–æ–≤–∫–æ–π; –ø—Ä–∏ –Ω–∞–ª–∏—á–∏–∏ ‚Äî `–ë–∞–Ω–∫–æ–≤—Å–∫–∏–µ —Ä–µ–∫–≤–∏–∑–∏—Ç—ã`, `–°–ø–æ—Å–æ–± –ø–æ–¥–∞—á–∏`). –ù–µ –∏—Å–ø–æ–ª—å–∑—É–π –ø—Å–µ–≤–¥–æ–≥—Ä–∞—Ñ–∏–∫—É –∏ –æ–±—Ä–∞—Ç–Ω—ã–µ —Å–ª—ç—à–∏.
   ‚Ä¢ –ü–æ—Å–ª–µ —Ä–µ–∫–≤–∏–∑–∏—Ç–æ–≤ –¥–æ–±–∞–≤—å –ø—É—Å—Ç—É—é —Å—Ç—Ä–æ–∫—É –∏ –∑–∞–≥–æ–ª–æ–≤–æ–∫ `# {–Ω–∞–∑–≤–∞–Ω–∏–µ}`.
   ‚Ä¢ –û—Å–Ω–æ–≤–Ω—ã–µ —Ä–∞–∑–¥–µ–ª—ã –æ—Ñ–æ—Ä–º–ª—è–π –ø–æ–¥–∑–∞–≥–æ–ª–æ–≤–∫–∞–º–∏ `## ‚Ä¶`; –∏—Å–ø–æ–ª—å–∑—É–π –Ω—É–º–µ—Ä–æ–≤–∞–Ω–Ω—ã–µ —Å–ø–∏—Å–∫–∏ `1.`/`1)` –¥–ª—è –∫–ª—é—á–µ–≤—ã—Ö –ø—É–Ω–∫—Ç–æ–≤ –∏ –º–∞—Ä–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ —Å–ø–∏—Å–∫–∏ `-`, `*`, `‚Ä¢`, `‚Äî` –¥–ª—è –ø–æ–¥–ø—É–Ω–∫—Ç–æ–≤.
   ‚Ä¢ –ò—Ç–æ–≥–æ–≤—ã–µ –±–ª–æ–∫–∏ (–ø—Ä–æ—Å–∏—Ç–µ–ª—å–Ω–∞—è —á–∞—Å—Ç—å, –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è, –ø–æ–¥–ø–∏—Å—å) –ø–æ–¥–∞–≤–∞–π –∫–∞–∫ —Å—Ç—Ä—É–∫—Ç—É—Ä–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ —Å–ø–∏—Å–∫–∏ –±–µ–∑ –æ–±—Ä–∞—Ç–Ω—ã—Ö —Å–ª—ç—à–µ–π –≤ –∫–æ–Ω—Ü–µ —Å—Ç—Ä–æ–∫.
   ‚Ä¢ –î–ª—è —Ä–∞—Å—á—ë—Ç–æ–≤ –∏ —Å–º–µ—Ç –∏—Å–ø–æ–ª—å–∑—É–π Markdown-—Ç–∞–±–ª–∏—Ü—ã (`| –∫–æ–ª–æ–Ω–∫–∞ | –∫–æ–ª–æ–Ω–∫–∞ |`), —á—Ç–æ–±—ã –∏—Ö –º–æ–∂–Ω–æ –±—ã–ª–æ –ø–µ—Ä–µ–Ω–µ—Å—Ç–∏ –≤ Word.
5. –ü–µ—Ä–µ–¥ –≤—ã–¥–∞—á–µ–π —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –≤—ã–ø–æ–ª–Ω–∏ —Å–∞–º–æ–ø—Ä–æ–≤–µ—Ä–∫—É: –ø–µ—Ä–µ—á–∏—Å–ª–∏ —É—á—Ç—ë–Ω–Ω—ã–µ –∫–ª—é—á–µ–≤—ã–µ –¥–∞–Ω–Ω—ã–µ –∏ –≤—ã—è–≤–ª–µ–Ω–Ω—ã–µ —Ä–∏—Å–∫–∏/–ø—Ä–æ–±–µ–ª—ã –≤ self_check.
6. –§–æ—Ä–º–∞—Ç –æ—Ç–≤–µ—Ç–∞:
{
  "status": "ok" | "need_more_info" | "abort",
  "document_title": "...",
  "document_markdown": "...",
  "self_check": {
    "validated": ["..."],
    "issues": ["..."]
  },
  "follow_up_questions": ["..."]
}
–ï—Å–ª–∏ status = "ok", –ø–æ–ª–µ document_markdown –¥–æ–ª–∂–Ω–æ —Å–æ–¥–µ—Ä–∂–∞—Ç—å –ø–æ–ª–Ω—ã–π —Ç–µ–∫—Å—Ç, –≥–æ—Ç–æ–≤—ã–π –∫ –∫–æ–Ω–≤–µ—Ä—Ç–∞—Ü–∏–∏ –≤ Word.
""".strip()

USER_TEMPLATE = """
{request_heading}
{request}

{details}
""".strip()


# ================================ Regexes =====================================

_JSON_START_RE = re.compile(r"[{\[]")
_TRAILING_COMMA_RE = re.compile(r",(\s*[}\]])")
_JSON_DECODER = json.JSONDecoder()
_INLINE_MARKUP_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`)", re.DOTALL)
_BULLET_ITEM_RE = re.compile(r"^([-*‚Ä¢‚Äî‚Äì])\s+(.*)")
_NUMBERED_ITEM_RE = re.compile(r"^\d+[.)]\s+(.*)")
_METADATA_LINE_RE = re.compile(r"^([^:]{1,80}):\s+(.+)$")
_TABLE_SEPARATOR_RE = re.compile(r":?-{3,}:?")


# =============================== JSON helpers =================================

def _strip_code_fences(payload: str) -> str:
    stripped = payload.strip()
    lowered = stripped.lower()
    for prefix in ("```json", "```", "~~~json", "~~~"):
        if lowered.startswith(prefix):
            stripped = stripped[len(prefix):]
            stripped = stripped.lstrip("\r\n")
            break

    for suffix in ("```", "~~~"):
        if stripped.endswith(suffix):
            stripped = stripped[:-len(suffix)]
            stripped = stripped.rstrip()
            break

    return stripped.strip()


def _sanitize_json_string(payload: str) -> str:
    """
    Escape literal newlines and stray double quotes that appear inside JSON strings.

    This function walks the payload character-by-character, keeping track of whether
    we are inside a JSON string literal. When inside a string it converts raw CR/LF
    characters to their escaped versions and protects accidental quote characters
    (e.g. "–ú–æ–π –∞—Ä–±–∏—Ç—Ä") that the model may emit without escaping.
    """

    if ("\n" not in payload and "\r" not in payload and '"' not in payload):
        return payload

    result: list[str] = []
    in_string = False
    escaped = False
    index = 0
    length = len(payload)

    while index < length:
        char = payload[index]

        if in_string:
            if escaped:
                result.append(char)
                escaped = False
            else:
                if char == "\\":
                    result.append(char)
                    escaped = True
                elif char == '"':
                    # Determine whether this quote should terminate the string.
                    lookahead = index + 1
                    while lookahead < length and payload[lookahead] in {" ", "\t"}:
                        lookahead += 1
                    next_char = payload[lookahead] if lookahead < length else ""

                    if next_char in {"\r", "\n"}:
                        while lookahead < length and payload[lookahead] in {" ", "\t", "\r", "\n"}:
                            lookahead += 1
                        next_char = payload[lookahead] if lookahead < length else ""

                    if next_char == ",":
                        lookahead += 1
                        while lookahead < length and payload[lookahead] in {" ", "\t", "\r", "\n"}:
                            lookahead += 1
                        after_comma = payload[lookahead] if lookahead < length else ""
                        if after_comma in {'"', "}", "]"}:
                            result.append(char)
                            in_string = False
                        else:
                            result.append('\\"')
                    elif next_char in {"", "}", "]", ":"}:
                        result.append(char)
                        in_string = False
                    else:
                        result.append('\\"')
                elif char == "\r":
                    result.append("\\r")
                    if index + 1 < length and payload[index + 1] == "\n":
                        result.append("\\n")
                        index += 1
                elif char == "\n":
                    result.append("\\n")
                else:
                    result.append(char)
        else:
            result.append(char)
            if char == '"':
                in_string = True
                escaped = False

        index += 1

    sanitized = "".join(result)
    return sanitized


def _deduplicate_consecutive_properties(payload: str) -> str:
    """Remove consecutive duplicate object properties that often appear due to streaming glitches."""

    lines = payload.splitlines()
    if not lines:
        return payload

    result: list[str] = []
    prev_key: str | None = None
    prev_value: str | None = None

    for line in lines:
        stripped = line.strip()
        normalized_line = stripped.lstrip(",")

        if normalized_line.startswith('"') and ":" in normalized_line:
            key_part, value_part = normalized_line.split(":", 1)
            key = key_part.strip().strip('"').strip()
            value = value_part.strip()
            value_no_comma = value[:-1].rstrip() if value.endswith(",") else value

            if prev_key == key:
                if prev_value != value_no_comma:
                    logger.debug(
                        "Dropping duplicate key with differing value in JSON repair: %s",
                        key,
                    )
                # Skip exact duplicate of the previous property entry
                continue

            prev_key = key
            prev_value = value_no_comma
        else:
            prev_key = None
            prev_value = None

        result.append(line)

    return "\n".join(result)


def _extract_structured_payload(raw: Any) -> Mapping[str, Any] | None:
    """Attempt to locate the actual JSON payload within structured response metadata."""

    def _unwrap(candidate: Any) -> Mapping[str, Any] | None:
        if isinstance(candidate, Mapping):
            lower_keys = {str(key).lower() for key in candidate.keys()}
            if (
                "document_title" in lower_keys
                or "document_markdown" in lower_keys
                or "status" in lower_keys
                or "questions" in lower_keys
                or "context_notes" in lower_keys
            ):
                return dict(candidate)
            nested_keys = (
                "parsed",
                "data",
                "json_schema",
                "content",
                "output",
                "outputs",
                "message",
                "messages",
                "result",
                "response",
                "payload",
                "body",
                "value",
            )
            for key in nested_keys:
                if key in candidate:
                    nested = candidate.get(key)
                    unwrapped = _unwrap(nested)
                    if unwrapped is not None:
                        return unwrapped
            for nested_value in candidate.values():
                if isinstance(nested_value, (Mapping, list, tuple)):
                    unwrapped = _unwrap(nested_value)
                    if unwrapped is not None:
                        return unwrapped
            return None

        if isinstance(candidate, (list, tuple)):
            for item in candidate:
                unwrapped = _unwrap(item)
                if unwrapped is not None:
                    return unwrapped
        return None

    return _unwrap(raw)


def _extract_json(text: Any) -> Any:
    """Extract the first JSON structure found in text."""
    if isinstance(text, Mapping):
        return text
    if isinstance(text, list):
        for item in text:
            if isinstance(item, Mapping):
                return item
        if text:
            return text[0]
        raise DocumentDraftingError("–ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞–π—Ç–∏ JSON –≤ –æ—Ç–≤–µ—Ç–µ –º–æ–¥–µ–ª–∏")
    if text is None:
        raise DocumentDraftingError("–ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞–π—Ç–∏ JSON –≤ –æ—Ç–≤–µ—Ç–µ –º–æ–¥–µ–ª–∏")
    if not isinstance(text, str):
        raise DocumentDraftingError("–ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞–π—Ç–∏ JSON –≤ –æ—Ç–≤–µ—Ç–µ –º–æ–¥–µ–ª–∏")

    cleaned_text = _strip_code_fences(text)

    def _decode_from_index(source: str, start_index: int = 0) -> Any:
        idx = start_index
        length = len(source)
        while idx < length and source[idx].isspace():
            idx += 1
        if idx >= length:
            raise json.JSONDecodeError("Empty JSON candidate", source, idx)
        try:
            obj, _ = _JSON_DECODER.raw_decode(source, idx)
            return obj
        except (json.JSONDecodeError, ValueError) as primary_err:
            candidate = source[idx:]
            stripped_candidate = candidate.strip()
            if not stripped_candidate:
                raise json.JSONDecodeError("Empty JSON candidate", source, idx) from primary_err

            normalized = _TRAILING_COMMA_RE.sub(r"\1", stripped_candidate)

            candidates: list[str] = [normalized]

            deduped = _deduplicate_consecutive_properties(normalized)
            if deduped != normalized:
                candidates.append(deduped)

            sanitized_variants: list[str] = []
            for candidate in candidates:
                sanitized = _sanitize_json_string(candidate)
                if sanitized != candidate:
                    sanitized_variants.append(sanitized)
            candidates.extend(sanitized_variants)

            for attempt in candidates:
                with suppress(json.JSONDecodeError, ValueError):
                    obj, _ = _JSON_DECODER.raw_decode(attempt)
                    return obj

            raise primary_err

    try:
        return _decode_from_index(cleaned_text)
    except (json.JSONDecodeError, ValueError) as err:
        logger.debug("Primary JSON parse failed, scanning for nested JSON: %s", err)
        for match in _JSON_START_RE.finditer(cleaned_text):
            start_idx = match.start()
            try:
                return _decode_from_index(cleaned_text, start_idx)
            except (json.JSONDecodeError, ValueError):
                continue

    raise DocumentDraftingError("–ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞–π—Ç–∏ JSON –≤ –æ—Ç–≤–µ—Ç–µ –º–æ–¥–µ–ª–∏")


def _looks_like_document_payload(data: Mapping[str, Any]) -> bool:
    status_value = str(data.get("status") or "").lower()
    if status_value in {"ok", "need_more_info", "abort"}:
        return True
    candidate_keys = {
        "document_title",
        "document_markdown",
        "questions",
        "follow_up_questions",
        "context_notes",
        "need_more_info",
        "self_check",
    }
    return any(key in data for key in candidate_keys)


def _coerce_payload(structured: Any, raw_text: str) -> Mapping[str, Any]:
    if isinstance(structured, Mapping):
        mapping = dict(structured)
        if _looks_like_document_payload(mapping):
            return mapping
        for key in ("parsed", "data", "json", "json_schema"):
            nested = mapping.get(key)
            if nested:
                try:
                    return _coerce_payload(nested, raw_text)
                except DocumentDraftingError:
                    continue
    if isinstance(structured, list):
        for item in structured:
            try:
                return _coerce_payload(item, raw_text)
            except DocumentDraftingError:
                continue
    if raw_text:
        extracted = _extract_json(raw_text)
        if isinstance(extracted, Mapping) and _looks_like_document_payload(extracted):
            return extracted
    if structured is not None:
        extracted = _extract_json(structured)
        if isinstance(extracted, Mapping) and _looks_like_document_payload(extracted):
            return extracted
    raise DocumentDraftingError("–ù–µ —É–¥–∞–ª–æ—Å—å –ø–æ–ª—É—á–∏—Ç—å —Å—Ç—Ä—É–∫—Ç—É—Ä–∏—Ä–æ–≤–∞–Ω–Ω—ã–π –æ—Ç–≤–µ—Ç –º–æ–¥–µ–ª–∏")


def _format_answers(answers: Iterable[dict[str, str]]) -> str:
    lines = []
    for idx, item in enumerate(answers, 1):
        question = item.get("question", "").strip()
        answer = item.get("answer", "").strip()
        lines.append(f"{idx}. {question}\n   –û—Ç–≤–µ—Ç: {answer}")
    return "\n".join(lines) if lines else "(–û—Ç–≤–µ—Ç—ã –ø–æ–∫–∞ –Ω–µ –ø–æ–ª—É—á–µ–Ω—ã)"


# ============================ High-level API calls ============================

async def plan_document(openai_service, request_text: str) -> DraftPlan:
    if not request_text.strip():
        raise DocumentDraftingError("–ü—É—Å—Ç–æ–π –∑–∞–ø—Ä–æ—Å")

    planner_details = """
–°—Ñ–æ—Ä–º–∏—Ä—É–π —Å—Ç—Ä—É–∫—Ç—É—Ä—É –æ—Ç–≤–µ—Ç–∞:
- –ü–æ–ª–µ document_title ‚Äî —Ç–æ—á–Ω–æ–µ –Ω–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞.
- –ü–æ–ª–µ need_more_info = true, –µ—Å–ª–∏ —Ç—Ä–µ–±—É—é—Ç—Å—è –æ—Ç–≤–µ—Ç—ã —é—Ä–∏—Å—Ç–∞; false, –µ—Å–ª–∏ –¥–∞–Ω–Ω—ã—Ö –¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏.
- –ú–∞—Å—Å–∏–≤ questions: —ç–ª–µ–º–µ–Ω—Ç—ã {"id": "...", "text": "...", "purpose": "..."} —Å –ø–æ—è—Å–Ω–µ–Ω–∏–µ–º —Ü–µ–ª–∏ –≤–æ–ø—Ä–æ—Å–∞.
- –ú–∞—Å—Å–∏–≤ context_notes: –∫–æ—Ä–æ—Ç–∫–∏–µ –∑–∞–º–µ—Ç–∫–∏, –∫–æ—Ç–æ—Ä—ã–µ –ø–æ–º–æ–≥—É—Ç –ø—Ä–∏ –ø–æ–¥–≥–æ—Ç–æ–≤–∫–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞.

–û—Ç–≤–µ—Ç –≤–µ—Ä–Ω–∏ —Å—Ç—Ä–æ–≥–æ –≤ JSON.
""".strip()

    user_prompt = USER_TEMPLATE.format(
        request_heading="–ó–∞–ø—Ä–æ—Å —é—Ä–∏—Å—Ç–∞:",
        request=request_text.strip(),
        details=planner_details,
    )
    response = await openai_service.ask_legal(
        system_prompt=DOCUMENT_ASSISTANT_SYSTEM_PROMPT,
        user_text=user_prompt,
    )
    if not response.get("ok"):
        raise DocumentDraftingError(response.get("error") or "–ù–µ —É–¥–∞–ª–æ—Å—å –ø–æ–ª—É—á–∏—Ç—å –æ—Ç–≤–µ—Ç –æ—Ç –º–æ–¥–µ–ª–∏")

    raw_text = response.get("text", "")
    structured = _extract_structured_payload(response.get("structured"))
    if structured:
        logger.debug("Using structured planner payload (keys: %s)", list(structured.keys()))
    data = _coerce_payload(structured, raw_text)

    title = str(data.get("document_title") or "–î–æ–∫—É–º–µ–Ω—Ç").strip()
    notes = [str(note).strip() for note in data.get("context_notes") or [] if str(note).strip()]

    need_more_raw = data.get("need_more_info")
    if isinstance(need_more_raw, bool):
        need_more = need_more_raw
    elif isinstance(need_more_raw, str):
        lowered = need_more_raw.strip().lower()
        need_more = lowered in {"true", "yes", "1", "y", "–¥–∞"}
    elif need_more_raw is None:
        need_more = False
    else:
        need_more = bool(need_more_raw)

    questions: list[dict[str, str]] = []
    if need_more:
        raw_questions = data.get("questions") or []
        for idx, raw in enumerate(raw_questions):
            text = str(raw.get("text") or "").strip()
            if not text:
                continue
            purpose = str(raw.get("purpose") or "").strip()
            question_id = str(raw.get("id") or f"q{idx + 1}")
            questions.append({"id": question_id, "text": text, "purpose": purpose})

    return DraftPlan(title=title or "–î–æ–∫—É–º–µ–Ω—Ç", questions=questions, notes=notes)


async def generate_document(
    openai_service,
    request_text: str,
    title: str,
    answers: list[dict[str, str]],
) -> DraftResult:
    answers_formatted = _format_answers(answers)
    generator_details = f"""
–ü—Ä–µ–¥–ø–æ–ª–∞–≥–∞–µ–º–æ–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {title}

–û—Ç–≤–µ—Ç—ã —é—Ä–∏—Å—Ç–∞:
{answers_formatted or "(–û—Ç–≤–µ—Ç—ã –ø–æ–∫–∞ –Ω–µ –ø–æ–ª—É—á–µ–Ω—ã)"}

–°–æ–±–µ—Ä–∏ –∏—Ç–æ–≥ –ø–æ —Å—Ö–µ–º–µ JSON:
{{
  "status": "ok" | "need_more_info" | "abort",
  "document_title": "...",
  "document_markdown": "...",
  "self_check": {{
    "validated": ["..."],
    "issues": ["..."]
  }},
  "follow_up_questions": ["..."]
}}
""".strip()
    user_prompt = USER_TEMPLATE.format(
        request_heading="–ó–∞–ø—Ä–æ—Å —é—Ä–∏—Å—Ç–∞ –∏ –≤–≤–æ–¥–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ:",
        request=request_text.strip(),
        details=generator_details,
    )
    response = await openai_service.ask_legal(
        system_prompt=DOCUMENT_ASSISTANT_SYSTEM_PROMPT,
        user_text=user_prompt,
    )
    if not response.get("ok"):
        raise DocumentDraftingError(response.get("error") or "–ù–µ —É–¥–∞–ª–æ—Å—å –ø–æ–ª—É—á–∏—Ç—å –æ—Ç–≤–µ—Ç –æ—Ç –º–æ–¥–µ–ª–∏")

    raw_text = response.get("text", "")
    structured = _extract_structured_payload(response.get("structured"))
    if structured:
        logger.debug("Using structured generator payload (keys: %s)", list(structured.keys()))
    data = _coerce_payload(structured, raw_text)

    status = str(data.get("status") or "ok").lower()
    doc_title = str(data.get("document_title") or title).strip()
    markdown = str(data.get("document_markdown") or "")
    self_check = data.get("self_check") or {}
    validated = [str(item).strip() for item in self_check.get("validated") or [] if str(item).strip()]
    issues = [str(item).strip() for item in self_check.get("issues") or [] if str(item).strip()]
    follow_up = [str(item).strip() for item in data.get("follow_up_questions") or [] if str(item).strip()]

    return DraftResult(
        status=status,
        title=doc_title or title,
        markdown=markdown,
        validated=validated,
        issues=issues,
        follow_up_questions=follow_up,
    )


# ============================= DOCX builder (MD) ==============================

def build_docx_from_markdown(markdown: str, output_path: str) -> None:
    """
    –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ—Ç –ø–æ–¥–≥–æ—Ç–æ–≤–ª–µ–Ω–Ω—ã–π Markdown –≤ DOCX, –ø—Ä–∏–º–µ–Ω—è—è –±–∞–∑–æ–≤—ã–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ —é—Ä–¥–æ–∫–∞–º:
    - –ë—É–º–∞–≥–∞ A4, –ø–æ–ª—è 2.5 —Å–º, —à—Ä–∏—Ñ—Ç Times New Roman 12 pt, –º–µ–∂–¥—É—Å—Ç—Ä–æ—á–Ω—ã–π –∏–Ω—Ç–µ—Ä–≤–∞–ª 1.5
    - –ó–∞–≥–æ–ª–æ–≤–∫–∏ —á–µ—Ä–µ–∑ —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–µ —Å—Ç–∏–ª–∏, –±–µ–∑ –ø—Å–µ–≤–¥–æ–≥—Ä–∞—Ñ–∏–∫–∏
    - –ú–µ—Ç–∞–¥–∞–Ω–Ω—ã–µ "–ü–æ–ª–µ: –ó–Ω–∞—á–µ–Ω–∏–µ" ‚Äî –∞–≤—Ç–æ–º–∞—Ç–æ–º —Å–∫–ª–∞–¥—ã–≤–∞—é—Ç—Å—è –≤ 2‚Äë–∫–æ–ª–æ–Ω–æ—á–Ω—É—é —Ç–∞–±–ª–∏—Ü—É
    - –ü–æ–¥–¥–µ—Ä–∂–∫–∞ —Å–ø–∏—Å–∫–æ–≤ (-, *, ‚Ä¢, ‚Äî) –∏ –Ω—É–º–µ—Ä–∞—Ü–∏–∏ (1. / 1))
    - –ü–æ–¥–¥–µ—Ä–∂–∫–∞ Markdown‚Äë—Ç–∞–±–ª–∏—Ü
    - –ù—É–º–µ—Ä–∞—Ü–∏—è —Å—Ç—Ä–∞–Ω–∏—Ü ¬´–°—Ç—Ä. X –∏–∑ Y¬ª –≤ –Ω–∏–∂–Ω–µ–º –∫–æ–ª–æ–Ω—Ç–∏—Ç—É–ª–µ
    """
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Cm, Pt  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise DocumentDraftingError("–ù–µ —É–¥–∞–ª–æ—Å—å –ø–æ–¥–≥–æ—Ç–æ–≤–∏—Ç—å DOCX –±–µ–∑ –ø–∞–∫–µ—Ç–∞ python-docx") from exc

    if not markdown or not markdown.strip():
        raise DocumentDraftingError("–ü—É—Å—Ç–æ–π —Ç–µ–∫—Å—Ç –¥–æ–∫—É–º–µ–Ω—Ç–∞")

    document = Document()
    section = document.sections[0]

    # –§–æ—Ä–º–∞—Ç —Å—Ç—Ä–∞–Ω–∏—Ü—ã –∏ –ø–æ–ª—è (A4 + 2.5 —Å–º)
    with suppress(Exception):
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.5))

    # ‚Äî‚Äî‚Äî —à—Ä–∏—Ñ—Ç—ã –∏ —Å—Ç–∏–ª–∏
    def _ensure_font(style_name: str, *, size: int, bold: bool = False,
                     italic: bool = False, align: int | None = None) -> None:
        try:
            style = document.styles[style_name]
        except KeyError:
            return
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(size)
        font.bold = bold
        font.italic = italic
        # –¥–ª—è –∫–∏—Ä–∏–ª–ª–∏—Ü—ã
        if style.element.rPr is not None:
            r_fonts = style.element.rPr.rFonts
            if r_fonts is not None:
                r_fonts.set(qn("w:eastAsia"), "Times New Roman")
        paragraph_format = getattr(style, "paragraph_format", None)
        if paragraph_format:
            paragraph_format.space_after = Pt(6)
            if align is not None:
                paragraph_format.alignment = align

    # Normal
    normal_style = document.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Times New Roman"
    normal_font.size = Pt(12)
    if normal_style.element.rPr is not None:
        r_fonts = normal_style.element.rPr.rFonts
        if r_fonts is not None:
            r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_paragraph = normal_style.paragraph_format
    normal_paragraph.space_after = Pt(6)
    normal_paragraph.first_line_indent = Cm(0)
    normal_paragraph.line_spacing = 1.5

    _ensure_font("Title", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _ensure_font("Heading 1", size=14, bold=True)
    _ensure_font("Heading 2", size=13, bold=True)
    _ensure_font("Heading 3", size=12, bold=True)

    # ‚Äî‚Äî‚Äî –∫–æ–ª–æ–Ω—Ç–∏—Ç—É–ª: "–°—Ç—Ä. X –∏–∑ Y"
    def _add_field(paragraph, instruction: str) -> None:
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_separate)
        run._r.append(fld_end)

    with suppress(Exception):
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.add_run("–°—Ç—Ä. ")
        _add_field(para, "PAGE")
        para.add_run(" –∏–∑ ")
        _add_field(para, "NUMPAGES")

    # ‚Äî‚Äî‚Äî inline —Ä–∞–∑–º–µ—Ç–∫–∞
    def _add_runs(paragraph, text: str) -> None:
        for chunk in _INLINE_MARKUP_RE.split(text):
            if not chunk:
                continue
            run = paragraph.add_run()
            if chunk.startswith(("**", "__")) and chunk.endswith(chunk[:2]):
                run.text = chunk[2:-2]
                run.bold = True
            elif chunk.startswith(("*", "_")) and chunk.endswith(chunk[0]):
                run.text = chunk[1:-1]
                run.italic = True
            elif chunk.startswith("`") and chunk.endswith("`"):
                run.text = chunk[1:-1]
                run.font.name = "Consolas"
                run.font.size = Pt(11)
            else:
                run.text = chunk

    # ‚Äî‚Äî‚Äî meta‚Äë–±–ª–æ–∫ —à–∞–ø–∫–∏
    list_mode: str | None = None
    metadata_buffer: list[tuple[str, str]] = []

    def _is_table_row(text: str) -> bool:
        stripped = text.strip()
        return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2

    def _parse_table_row(text: str) -> list[str]:
        cells = [cell.strip() for cell in text.strip()[1:-1].split("|")]
        return cells

    def flush_metadata(force_plain: bool = False) -> None:
        nonlocal metadata_buffer
        if not metadata_buffer:
            return

        if not force_plain and len(metadata_buffer) >= 2:
            table = document.add_table(rows=0, cols=2)
            with suppress(Exception):
                table.style = "Table Grid"
            for field, value in metadata_buffer:
                row = table.add_row()
                field_cell, value_cell = row.cells

                fp = field_cell.paragraphs[0]
                fp.paragraph_format.space_before = Pt(0)
                fp.paragraph_format.space_after = Pt(0)
                fp.paragraph_format.first_line_indent = Cm(0)
                fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fr = fp.add_run(field)
                fr.bold = True

                vp = value_cell.paragraphs[0]
                vp.paragraph_format.space_before = Pt(0)
                vp.paragraph_format.space_after = Pt(0)
                vp.paragraph_format.first_line_indent = Cm(0)
                vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _add_runs(vp, value)

            document.add_paragraph("")  # —Ä–∞–∑–¥–µ–ª–µ–Ω–∏–µ –æ—Ç –æ—Å—Ç–∞–ª—å–Ω–æ–≥–æ —Ç–µ–∫—Å—Ç–∞
        else:
            for field, value in metadata_buffer:
                paragraph = document.add_paragraph(style="Normal")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.left_indent = Cm(0)
                paragraph.paragraph_format.space_after = Pt(6)
                _add_runs(paragraph, f"{field}: {value}")

        metadata_buffer = []

    # ‚Äî‚Äî‚Äî –æ—Å–Ω–æ–≤–Ω–æ–π —Ä–∞–∑–±–æ—Ä Markdown
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").splitlines()
    idx = 0
    total_lines = len(lines)

    while idx < total_lines:
        raw_line = lines[idx]
        idx += 1
        line = raw_line.rstrip()

        # –ø–æ–¥–¥–µ—Ä–∂–∫–∞ "–∂—ë—Å—Ç–∫–æ–≥–æ –ø–µ—Ä–µ–Ω–æ—Å–∞" –≤ Markdown (backslash –≤ –∫–æ–Ω—Ü–µ —Å—Ç—Ä–æ–∫–∏)
        if line.endswith("\\"):
            line = line[:-1].rstrip()

        if not line:
            flush_metadata()
            document.add_paragraph("")
            list_mode = None
            continue

        plain_line = line.strip()
        plain_lower = plain_line.lower()

        # —ç–≤—Ä–∏—Å—Ç–∏–∫–∞ –¥–ª—è "–ò—Å–∫–æ–≤–æ–µ –∑–∞—è–≤–ª–µ–Ω–∏–µ ..." –≤ –ø–µ—Ä–≤–æ–π —Å—Ç—Ä–æ–∫–µ ‚Äî –∫–∞–∫ —Ç–∏—Ç—É–ª
        if plain_line and plain_lower.startswith("–∏—Å–∫–æ–≤–æ–µ –∑–∞—è–≤–ª–µ–Ω–∏–µ"):
            title_paragraph = document.add_paragraph(style="Title")
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_paragraph.paragraph_format.space_before = Pt(12)
            title_paragraph.paragraph_format.space_after = Pt(12)
            title_paragraph.paragraph_format.keep_with_next = True
            _add_runs(title_paragraph, plain_line)
            list_mode = None
            continue

        # —Å–±–æ—Ä –º–µ—Ç–∞‚Äë—Å—Ç—Ä–æ–∫ "–ü–æ–ª–µ: –ó–Ω–∞—á–µ–Ω–∏–µ" –¥–æ –ø–µ—Ä–≤–æ–≥–æ "–æ–±—ã—á–Ω–æ–≥–æ" –∞–±–∑–∞—Ü–∞/–∑–∞–≥–æ–ª–æ–≤–∫–∞/—Å–ø–∏—Å–∫–∞
        if list_mode is None and plain_line:
            if not plain_line.startswith(("-", "*", "‚Ä¢", "‚Äî", "#")) and not _NUMBERED_ITEM_RE.match(plain_line):
                colon_match = _METADATA_LINE_RE.match(plain_line)
                if colon_match:
                    field = colon_match.group(1).strip()
                    value = colon_match.group(2).strip()
                    if value:
                        metadata_buffer.append((field, value))
                        continue

        # —Ç–∞–±–ª–∏—Ü–∞ –≤ Markdown
        if _is_table_row(line):
            flush_metadata()
            table_lines = [line]
            while idx < total_lines:
                peek_line = lines[idx].rstrip()
                if _is_table_row(peek_line):
                    table_lines.append(peek_line)
                    idx += 1
                else:
                    break

            parsed_rows = [_parse_table_row(row) for row in table_lines]
            # –≤—Ç–æ—Ä–∞—è —Å—Ç—Ä–æ–∫–∞ ‚Äî —Ä–∞–∑–¥–µ–ª–∏—Ç–µ–ª—å?
            if len(parsed_rows) >= 2 and all(_TABLE_SEPARATOR_RE.fullmatch(cell.replace(" ", "")) for cell in parsed_rows[1]):
                data_rows = [parsed_rows[0]] + parsed_rows[2:]
            else:
                data_rows = parsed_rows

            if data_rows:
                num_cols = max(len(row) for row in data_rows)
                table = document.add_table(rows=0, cols=num_cols)
                with suppress(Exception):
                    table.style = "Table Grid"
                for row_idx, cells in enumerate(data_rows):
                    row = table.add_row()
                    for col in range(num_cols):
                        cell_text = cells[col].strip() if col < len(cells) else ""
                        cell = row.cells[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.first_line_indent = Cm(0)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        _add_runs(paragraph, cell_text)
                        if row_idx == 0:
                            for run in paragraph.runs:
                                run.bold = True
                document.add_paragraph("")
            list_mode = None
            continue

        flush_metadata()

        # –∑–∞–≥–æ–ª–æ–≤–∫–∏
        if line.startswith("# "):
            content = line[2:].strip()
            title_paragraph = document.add_paragraph(style="Title")
            title_paragraph.paragraph_format.space_after = Pt(12)
            title_paragraph.paragraph_format.keep_with_next = True
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(title_paragraph, content)  # –±–µ–∑ –Ω–∞—Å–∏–ª—å–Ω–æ–≥–æ UPPERCASE ‚Äî —é—Ä–∏–¥–∏—á–µ—Å–∫–∏–π —Å—Ç–∏–ª—å
            list_mode = None
            continue

        if line.startswith("## "):
            content = line[3:].strip()
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(12)
            _add_runs(paragraph, content)
            list_mode = None
            continue

        if line.startswith("### "):
            content = line[4:].strip()
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(10)
            _add_runs(paragraph, content)
            list_mode = None
            continue

        # —Å–ø–∏—Å–∫–∏
        bullet_match = _BULLET_ITEM_RE.match(line)
        number_match = _NUMBERED_ITEM_RE.match(line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Cm(0.75)
            _add_runs(paragraph, bullet_match.group(2).strip())
            list_mode = "bullet"
            continue
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Cm(0.75)
            _add_runs(paragraph, number_match.group(1).strip())
            list_mode = "number"
            continue

        # –æ–±—ã—á–Ω—ã–µ –∞–±–∑–∞—Ü—ã
        paragraph = document.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(6)
        if list_mode:
            paragraph.paragraph_format.left_indent = Cm(1.25)
            paragraph.paragraph_format.first_line_indent = Cm(0)
        else:
            # —Å—Ç—Ä–æ–∫–∞ –≤–∏–¥–∞ "–ü–æ–ª–µ: –∑–Ω–∞—á–µ–Ω–∏–µ" ‚Äî –±–µ–∑ –∫—Ä–∞—Å–Ω–æ–π —Å—Ç—Ä–æ–∫–∏, —Å–ª–µ–≤–∞
            if ":" in plain_line and not plain_line.endswith(":"):
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.left_indent = Cm(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.paragraph_format.first_line_indent = Cm(1.25)
        _add_runs(paragraph, plain_line)
        list_mode = None

    flush_metadata()

    document.save(output_path)


# ================================ UI helpers ==================================

def _pluralize_questions(count: int) -> str:
    """–°–∫–ª–æ–Ω–µ–Ω–∏–µ —Å–ª–æ–≤–∞ '–≤–æ–ø—Ä–æ—Å' –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç —á–∏—Å–ª–∞."""
    if count % 10 == 1 and count % 100 != 11:
        return "–≤–æ–ø—Ä–æ—Å"
    elif count % 10 in (2, 3, 4) and count % 100 not in (12, 13, 14):
        return "–≤–æ–ø—Ä–æ—Å–∞"
    else:
        return "–≤–æ–ø—Ä–æ—Å–æ–≤"


def format_plan_summary(plan: DraftPlan) -> str:
    from html import escape as html_escape

    lines: list[str] = []

    title = (plan.title or "–î–æ–∫—É–º–µ–Ω—Ç").strip()
    title_escaped = html_escape(title)

    # –ö—Ä–∞—Å–∏–≤—ã–π –∑–∞–≥–æ–ª–æ–≤–æ–∫ –¥–æ–∫—É–º–µ–Ω—Ç–∞
    lines.append("üìã <b>–ü–æ–¥–≥–æ—Ç–æ–≤–∫–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞</b>")
    lines.append("‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ‚îÅ")
    lines.append("")
    lines.append(f"<b>–î–æ–∫—É–º–µ–Ω—Ç:</b> {title_escaped}")
    lines.append("")

    if plan.questions:
        lines.append(f"<b>–¢—Ä–µ–±—É–µ—Ç—Å—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è:</b> {len(plan.questions)} {_pluralize_questions(len(plan.questions))}")
        lines.append("")
        lines.append("")
        lines.append("üí° <b>–ò–Ω—Å—Ç—Ä—É–∫—Ü–∏—è –ø–æ –æ—Ç–≤–µ—Ç—É</b>")
        lines.append("")
        lines.append("<b>–ö–∞–∫ –æ—Ç–≤–µ—á–∞—Ç—å:</b>")
        lines.append("")
        lines.append("  ‚úì –ù–∞–ø–∏—à–∏—Ç–µ –≤—Å–µ –æ—Ç–≤–µ—Ç—ã –≤ –æ–¥–Ω–æ–º")
        lines.append("    —Å–æ–æ–±—â–µ–Ω–∏–∏, –Ω–µ —Ä–∞–∑–¥–µ–ª—è—è –∏—Ö")
        lines.append("")
        lines.append("<b>–§–æ—Ä–º–∞—Ç –æ—Ç–≤–µ—Ç–æ–≤:</b>")
        lines.append("")
        lines.append("  <b>–í–∞—Ä–∏–∞–Ω—Ç 1</b> ‚Äî –Ω—É–º–µ—Ä–∞—Ü–∏—è:")
        lines.append("  <code>1) –í–∞—à –ø–µ—Ä–≤—ã–π –æ—Ç–≤–µ—Ç</code>")
        lines.append("  <code>2) –í–∞—à –≤—Ç–æ—Ä–æ–π –æ—Ç–≤–µ—Ç</code>")
        lines.append("  <code>3) –í–∞—à —Ç—Ä–µ—Ç–∏–π –æ—Ç–≤–µ—Ç</code>")
        lines.append("")
        lines.append("  <b>–í–∞—Ä–∏–∞–Ω—Ç 2</b> ‚Äî –ø—É—Å—Ç–∞—è —Å—Ç—Ä–æ–∫–∞:")
        lines.append("  <code>–ü–µ—Ä–≤—ã–π –æ—Ç–≤–µ—Ç</code>")
        lines.append("  <code></code>")
        lines.append("  <code>–í—Ç–æ—Ä–æ–π –æ—Ç–≤–µ—Ç</code>")
        lines.append("")
        lines.append("")
        lines.append("üëá <i>–í–æ–ø—Ä–æ—Å—ã –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω—ã –æ—Ç–¥–µ–ª—å–Ω—ã–º</i>")
        lines.append("   <i>—Å–æ–æ–±—â–µ–Ω–∏–µ–º –Ω–∏–∂–µ</i>")
    else:
        lines.append("<b>–°—Ç–∞—Ç—É—Å:</b> –ì–æ—Ç–æ–≤ –∫ —Å–æ–∑–¥–∞–Ω–∏—é ‚úÖ")
        lines.append("")
        lines.append("")
        lines.append("‚úÖ <b>–í—Å—ë –≥–æ—Ç–æ–≤–æ!</b>")
        lines.append("")
        lines.append("–î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è")
        lines.append("–Ω–µ —Ç—Ä–µ–±—É–µ—Ç—Å—è.")
        lines.append("")
        lines.append("üöÄ –ú–æ–∂–Ω–æ –ø—Ä–∏—Å—Ç—É–ø–∞—Ç—å –∫")
        lines.append("   —Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—é –¥–æ–∫—É–º–µ–Ω—Ç–∞")

    return "\n".join(lines)
