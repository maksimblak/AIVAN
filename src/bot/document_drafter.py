"""Utilities for interactive legal document drafting."""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass
from typing import Any, Iterable

logger = logging.getLogger(__name__)


class DocumentDraftingError(RuntimeError):
    """Raised when the LLM response cannot be processed."""


@dataclass(slots=True)
class DraftPlan:
    title: str
    questions: list[dict[str, str]]
    notes: list[str]


@dataclass(slots=True)
class DraftResult:
    status: str
    title: str
    markdown: str
    validated: list[str]
    issues: list[str]
    follow_up_questions: list[str]


DOCUMENT_PLANNER_SYSTEM_PROMPT = """
Ты — юридический ассистент, который помогает юристу подготовить точный правовой документ под конкретное дело.
Работай строго по этапам:

1. Если юрист уже назвал тип документа — найди подходящий образец и сформулируй точное название документа. Если тип не очевиден, задай наводящие вопросы, чтобы определить, какой именно документ нужен.
2. Подготовь перечень вопросов, которые необходимо задать юристу для заполнения документа. При необходимости группируй их по этапам и обязательно поясняй цель каждого вопроса.
3. Сформируй заметки контекста (context_notes), которые пригодятся при генерации документа и помогут ИИ придерживаться юридической точности.

Всегда возвращай валидный JSON.
""".strip()

DOCUMENT_PLANNER_USER_TEMPLATE = """
Запрос юриста:
{request}

Сформируй структуру ответа:
- Поле document_title — точное наименование документа.
- Поле need_more_info = true, если требуются ответы юриста; false, если данных достаточно для генерации.
- Массив questions: элементы {{"id": "...", "text": "...", "purpose": "..."}} с пояснением цели вопроса.
- Массив context_notes: короткие заметки, которые помогут при подготовке документа.

Ответ верни строго в JSON.
""".strip()

DOCUMENT_GENERATOR_SYSTEM_PROMPT = """
Ты — юридический ассистент-документалист. На основе запроса и ответов юриста подготовь юридически точный документ.
Следуй инструкции:

1. Проверь, достаточно ли данных. Если чего-то не хватает, верни status = "need_more_info" и сформируй follow_up_questions с уточнениями.
2. Если данных достаточно, составь документ в Markdown, опираясь на найденный образец и профессиональные стандарты. Соблюдай структуру, стиль и формулировки, характерные для соответствующего вида документа.
3. Перед выдачей результата выполни самопроверку: перечисли учтённые ключевые данные и выявленные риски/пробелы.

Формат ответа — валидный JSON вида:
{
  "status": "ok" | "need_more_info" | "abort",
  "document_title": "...",
  "document_markdown": "...",
  "self_check": {
    "validated": ["..."],
    "issues": ["..."]
  },
  "follow_up_questions": ["..."]
}
Если status = "ok", поле document_markdown должно содержать полный текст, готовый к конвертации в Word.
""".strip()

DOCUMENT_GENERATOR_USER_TEMPLATE = """
Запрос юриста и вводные данные:
{request}

Предполагаемое название документа: {title}

Ответы юриста:
{answers}

Собери итог по схеме JSON:
{{
  "status": "ok" | "need_more_info" | "abort",
  "document_title": "...",
  "document_markdown": "...",
  "self_check": {{
    "validated": ["..."],
    "issues": ["..."]
  }},
  "follow_up_questions": ["..."]
}}
""".strip()

_JSON_RE = re.compile(r"\{[\s\S]*\}")
_JSON_ARRAY_RE = re.compile(r"\[[\s\S]*\]")
_TRAILING_COMMA_RE = re.compile(r",(\s*[}\]])")


def _extract_json(text: str) -> Any:
    """Extract the first JSON object found in text."""

    def _try_parse(candidate: str) -> Any:
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            cleaned = _TRAILING_COMMA_RE.sub(r"\1", candidate)
            return json.loads(cleaned)

    try:
        return _try_parse(text)
    except json.JSONDecodeError:
        for regex in (_JSON_RE, _JSON_ARRAY_RE):
            match = regex.search(text)
            if not match:
                continue
            try:
                return _try_parse(match.group(0))
            except json.JSONDecodeError:
                continue
        raise DocumentDraftingError("Не удалось найти JSON в ответе модели")


def _format_answers(answers: Iterable[dict[str, str]]) -> str:
    lines = []
    for idx, item in enumerate(answers, 1):
        question = item.get("question", "").strip()
        answer = item.get("answer", "").strip()
        lines.append(f"{idx}. {question}\n   Ответ: {answer}")
    return "\n".join(lines) if lines else "(Ответы пока не получены)"



async def plan_document(openai_service, request_text: str) -> DraftPlan:
    if not request_text.strip():
        raise DocumentDraftingError("Пустой запрос")

    user_prompt = DOCUMENT_PLANNER_USER_TEMPLATE.format(request=request_text.strip())
    response = await openai_service.ask_legal(
        system_prompt=DOCUMENT_PLANNER_SYSTEM_PROMPT,
        user_text=user_prompt,
    )
    if not response.get("ok"):
        raise DocumentDraftingError(response.get("error") or "Не удалось получить ответ от модели")

    raw_text = response.get("text", "")
    if not raw_text:
        raise DocumentDraftingError("Пустой ответ модели")

    data = _extract_json(raw_text)
    title = str(data.get("document_title") or "Документ").strip()
    questions = [
        {
            "id": str(q.get("id") or f"q{i+1}"),
            "text": str(q.get("text") or "").strip(),
            "purpose": str(q.get("purpose") or "").strip(),
        }
        for i, q in enumerate(data.get("questions") or [])
        if str(q.get("text") or "").strip()
    ]
    notes = [str(note).strip() for note in data.get("context_notes") or [] if str(note).strip()]

    need_more = bool(data.get("need_more_info"))
    if not need_more:
        questions = []

    return DraftPlan(title=title or "Документ", questions=questions, notes=notes)


async def generate_document(
    openai_service,
    request_text: str,
    title: str,
    answers: list[dict[str, str]],
) -> DraftResult:
    answers_formatted = _format_answers(answers)
    user_prompt = DOCUMENT_GENERATOR_USER_TEMPLATE.format(
        request=request_text.strip(),
        title=title,
        answers=answers_formatted or "(Ответы пока не получены)",
    )
    response = await openai_service.ask_legal(
        system_prompt=DOCUMENT_GENERATOR_SYSTEM_PROMPT,
        user_text=user_prompt,
    )
    if not response.get("ok"):
        raise DocumentDraftingError(response.get("error") or "Не удалось получить ответ от модели")

    raw_text = response.get("text", "")
    if not raw_text:
        raise DocumentDraftingError("Пустой ответ модели")

    data = _extract_json(raw_text)
    status = str(data.get("status") or "ok").lower()
    doc_title = str(data.get("document_title") or title).strip()
    markdown = str(data.get("document_markdown") or "")
    self_check = data.get("self_check") or {}
    validated = [str(item).strip() for item in self_check.get("validated") or [] if str(item).strip()]
    issues = [str(item).strip() for item in self_check.get("issues") or [] if str(item).strip()]
    follow_up = [str(item).strip() for item in data.get("follow_up_questions") or [] if str(item).strip()]

    return DraftResult(
        status=status,
        title=doc_title or title,
        markdown=markdown,
        validated=validated,
        issues=issues,
        follow_up_questions=follow_up,
    )


def build_docx_from_markdown(markdown: str, output_path: str) -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise DocumentDraftingError("Для генерации DOCX требуется python-docx") from exc

    if not markdown.strip():
        raise DocumentDraftingError("Пустой текст документа")

    document = Document()
    current_list = False

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            if current_list:
                current_list = False
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            current_list = False
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            current_list = False
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            current_list = False
        elif line.startswith(('- ', '* ')):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            current_list = True
        else:
            document.add_paragraph(line)
            current_list = False

    document.save(output_path)


def format_plan_summary(plan: DraftPlan) -> str:
    lines: list[str] = []

    title = (plan.title or "Документ").strip()
    lines.append(f"📄 Документ: {title}")
    lines.append("━━━━━━━━━━━━━━━━━━━━")

    if plan.notes:
        lines.append("🔍 Контекст (важное):")
        max_notes = 6
        for note in plan.notes[:max_notes]:
            clean = str(note).strip()
            if clean:
                lines.append(f"• {clean}")
        remaining = len(plan.notes) - max_notes
        if remaining > 0:
            lines.append(f"… ещё {remaining} пункт(ов)")
        lines.append("")

    lines.append(f"❓ Уточняющих вопросов: {len(plan.questions)}")
    if plan.questions:
        lines.append("💬 Ответьте на них одним сообщением, следуя подсказке ниже.")
    else:
        lines.append("✅ Дополнительных уточнений не требуется — можно формировать документ.")

    return "\n".join(lines)
